#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
智能知识库系统 (DeepSeek版)

Copyright (c) 2026 吕滢

Licensed under the MIT License (Non-Commercial) or Apache License 2.0 (Non-Commercial)
See LICENSE-MIT-NC or LICENSE-APACHE-NC for details.

This software is for NON-COMMERCIAL USE ONLY.
For commercial use, please contact the copyright holder.
"""

import streamlit as st
import os
import glob
from typing import List, Dict, Any, Optional
import tempfile
from pathlib import Path
import json
from datetime import datetime
import base64
import hashlib

# API Key 管理模块
CONFIG_FILE = os.path.join(".", ".deepseek_config.json")

# Prompt 模版管理模块
PROMPT_TEMPLATES_DIR = os.path.join(".", "prompt_templates")
SUMMARY_TEMPLATES_FILE = os.path.join(PROMPT_TEMPLATES_DIR, "summary_templates.json")
ANALYSIS_TEMPLATES_FILE = os.path.join(PROMPT_TEMPLATES_DIR, "analysis_templates.json")

def ensure_templates_dir():
    """确保模版目录存在"""
    os.makedirs(PROMPT_TEMPLATES_DIR, exist_ok=True)

def is_default_template(template_type: str, template_id: str) -> bool:
    """检查模版是否是默认模版（不可删除）
    
    Args:
        template_type: 'summary' 或 'analysis'
        template_id: 模版ID
    
    Returns:
        如果是默认模版返回True，否则返回False
    """
    if template_type == "summary":
        default_template_ids = ["default", "brief", "detailed", "roadmap", "gantt"]
        return template_id in default_template_ids
    elif template_type == "analysis":
        default_template_ids = ["default", "statistical", "trend"]
        return template_id in default_template_ids
    return False

def get_default_summary_templates() -> Dict[str, Dict[str, Any]]:
    """获取默认的总结模版"""
    return {
        "default": {
            "name": "默认总结模版",
            "description": "标准的多部分总结报告",
            "template": """请根据以下文档内容，生成一份详细的总结报告：

文档内容：
{content}

请生成包括以下部分的报告：
1. 整体内容概述
2. 核心要点总结
3. 关键数据/信息提取
4. 主要发现和洞察
5. 建议和下一步行动

报告：""",
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        },
        "brief": {
            "name": "简要总结模版",
            "description": "简洁的要点总结",
            "template": """请根据以下文档内容，生成一份简要总结：

文档内容：
{content}

请提供：
1. 核心要点（3-5条）
2. 关键信息
3. 主要结论

总结：""",
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        },
        "detailed": {
            "name": "详细分析模版",
            "description": "深入分析文档内容",
            "template": """请对以下文档内容进行深入分析：

文档内容：
{content}

请提供详细分析：
1. 文档背景和目的
2. 主要内容结构
3. 关键数据和事实
4. 深度洞察和分析
5. 潜在问题和风险
6. 改进建议和行动计划

分析报告：""",
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        },
        "roadmap": {
            "name": "技术路线图模版",
            "description": "生成符号化技术路线图",
            "template": """请根据所选文档内容，制作项目的技术路线图（用线条符号来尝试组件技术路线图）。要求：
· 主要阶段使用方括号 [] 包裹，通过向下箭头 ↓ 连接
· 子任务通过向右箭头 → 连接
· 支持缩进表示层级关系

文档内容：
{content}

技术路线图：""",
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        },
        "gantt": {
            "name": "项目进度甘特图模版",
            "description": "生成项目进度甘特图表数据（表格格式）",
            "template": """请根据以下文档内容，生成一份详细的项目进度甘特图表数据。

文档内容：
{content}

## 输出格式要求

请严格按照以下表格格式输出，使用制表符（Tab）分隔各列：

```
任务ID	任务名称	开始时间	结束时间	工期(月)	前置任务	责任方/备注
```

## 列说明

1. **任务ID**：任务的唯一标识符
   - 主要阶段：1, 2, 3, 4...
   - 子任务：1.1, 1.2, 2.1, 2.2...
   - 三级任务：1.1.1, 1.1.2...

2. **任务名称**：任务的描述名称

3. **开始时间**：使用 M+数字 格式
   - M0：项目开始（第0个月）
   - M1：第1个月
   - M1+0.5 或 M1.5：第1.5个月
   - 示例：M0, M0+0.5, M1, M1.5, M2

4. **结束时间**：使用 M+数字 格式（必须大于等于开始时间）

5. **工期(月)**：任务持续时间（可选，会自动计算）
   - 可以是小数：0.5, 1, 1.5, 2, 2.5, 3

6. **前置任务**：依赖的任务ID（可选）
   - 多个任务用逗号或空格分隔：1.1, 1.2 或 2.1 2.2 2.3
   - 如果无前置任务，留空

7. **责任方/备注**：任务的责任人或备注信息（可选）
   - 格式：责任方（备注说明）
   - 示例：乙方（输出《需求分析说明书》）
   - 示例：甲方、乙方
   - 如果无备注，留空

## 任务层级要求

- **主要阶段（level 0）**：任务ID为单个数字（1, 2, 3...）
- **二级任务（level 1）**：任务ID为 X.Y 格式（1.1, 1.2, 2.1...）
- **三级任务（level 2）**：任务ID为 X.Y.Z 格式（1.1.1, 1.1.2...）

## 时间规划要求

1. **时间连续性**：确保任务时间顺序合理，前置任务完成后才能开始后续任务
2. **时间重叠**：允许并行任务，但需明确标注前置依赖关系
3. **时间跨度**：根据项目实际情况设定
4. **里程碑**：主要阶段应设置明确的开始和结束时间

## 注意事项

1. **必须包含表头行**：第一行必须是列标题
2. **使用制表符分隔**：列之间使用Tab键分隔，不要使用空格
3. **时间格式统一**：统一使用 M+数字 格式
4. **任务ID唯一性**：确保每个任务ID唯一
5. **依赖关系正确**：前置任务ID必须存在于任务列表中
6. **层级结构清晰**：主要阶段、二级任务、三级任务层次分明
7. **备注信息完整**：尽量为每个任务提供责任方和备注信息

请根据文档内容中的项目信息，生成完整的项目进度甘特图表数据。""",
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    }

def get_default_analysis_templates() -> Dict[str, Dict[str, Any]]:
    """获取默认的数据分析模版"""
    return {
        "default": {
            "name": "默认分析模版",
            "description": "标准的数据分析报告",
            "template": """请分析以下文档集合，提供数据分析:

文档信息：
{doc_info}

请提供：
1. 文档内容分布分析
2. 潜在的数据模式和趋势
3. 建议的数据可视化方式

分析结果：""",
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        },
        "statistical": {
            "name": "统计分析模版",
            "description": "侧重于统计数据分析",
            "template": """请对以下文档集合进行统计分析:

文档信息：
{doc_info}

请提供：
1. 文档数量、大小、类型分布统计
2. 内容关键词频率分析
3. 文档间关联性分析
4. 数据质量评估
5. 统计图表建议

统计分析：""",
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        },
        "trend": {
            "name": "趋势分析模版",
            "description": "侧重于趋势和模式识别",
            "template": """请分析以下文档集合中的趋势和模式:

文档信息：
{doc_info}

请提供：
1. 内容趋势识别
2. 时间序列模式（如有）
3. 主题演变趋势
4. 异常模式检测
5. 未来趋势预测

趋势分析：""",
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    }

def load_templates(template_type: str) -> Dict[str, Dict[str, Any]]:
    """加载模版（总结或分析）
    
    Args:
        template_type: 'summary' 或 'analysis'
    
    Returns:
        模版字典
    """
    ensure_templates_dir()
    
    if template_type == "summary":
        file_path = SUMMARY_TEMPLATES_FILE
        default_templates = get_default_summary_templates()
    elif template_type == "analysis":
        file_path = ANALYSIS_TEMPLATES_FILE
        default_templates = get_default_analysis_templates()
    else:
        return {}
    
    try:
        if os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as f:
                templates = json.load(f)
                # 合并默认模版（如果用户模版中没有）
                for key, default_template in default_templates.items():
                    if key not in templates:
                        templates[key] = default_template
                return templates
        else:
            # 如果文件不存在，创建默认模版文件
            save_templates(template_type, default_templates)
            return default_templates
    except Exception:
        # 如果加载失败，返回默认模版
        return default_templates

def save_templates(template_type: str, templates: Dict[str, Dict[str, Any]]) -> bool:
    """保存模版
    
    Args:
        template_type: 'summary' 或 'analysis'
        templates: 模版字典
    
    Returns:
        是否保存成功
    """
    ensure_templates_dir()
    
    if template_type == "summary":
        file_path = SUMMARY_TEMPLATES_FILE
    elif template_type == "analysis":
        file_path = ANALYSIS_TEMPLATES_FILE
    else:
        return False
    
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(templates, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        if 'st' in globals():
            st.error(f"保存模版失败: {str(e)}")
        return False

def save_template(template_type: str, template_id: str, name: str, description: str, template: str) -> bool:
    """保存单个模版
    
    Args:
        template_type: 'summary' 或 'analysis'
        template_id: 模版ID（如果已存在则更新，否则创建）
        name: 模版名称
        description: 模版描述
        template: 模版内容
    
    Returns:
        是否保存成功
    """
    templates = load_templates(template_type)
    
    # 生成ID（如果未提供或已存在）
    if not template_id or template_id in templates:
        # 使用名称生成ID（移除特殊字符）
        template_id = "".join(c if c.isalnum() or c in ('-', '_') else '_' for c in name.lower())
        # 确保唯一性
        counter = 1
        original_id = template_id
        while template_id in templates:
            template_id = f"{original_id}_{counter}"
            counter += 1
    
    templates[template_id] = {
        "name": name,
        "description": description,
        "template": template,
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    
    return save_templates(template_type, templates)

def delete_template(template_type: str, template_id: str) -> bool:
    """删除模版（默认模版不可删除）
    
    Args:
        template_type: 'summary' 或 'analysis'
        template_id: 模版ID
    
    Returns:
        是否删除成功
    """
    # 检查是否是默认模版，默认模版不可删除
    if is_default_template(template_type, template_id):
        return False
    
    templates = load_templates(template_type)
    
    if template_id in templates:
        del templates[template_id]
        return save_templates(template_type, templates)
    
    return False

def get_template(template_type: str, template_id: str) -> Optional[Dict[str, Any]]:
    """获取单个模版
    
    Args:
        template_type: 'summary' 或 'analysis'
        template_id: 模版ID
    
    Returns:
        模版字典，如果不存在则返回None
    """
    templates = load_templates(template_type)
    return templates.get(template_id)

def encode_api_key(api_key: str) -> str:
    """简单的编码（Base64），不是真正的加密，但可以避免完全明文"""
    if not api_key:
        return ""
    # 使用 Base64 编码
    encoded = base64.b64encode(api_key.encode('utf-8')).decode('utf-8')
    return encoded

def decode_api_key(encoded_key: str) -> str:
    """解码 API key"""
    if not encoded_key:
        return ""
    try:
        decoded = base64.b64decode(encoded_key.encode('utf-8')).decode('utf-8')
        return decoded
    except Exception:
        return ""

def save_api_key(api_key: str, show_error: bool = True) -> bool:
    """保存 API key 到本地配置文件"""
    try:
        config = {
            "api_key": encode_api_key(api_key),
            "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        os.makedirs(os.path.dirname(CONFIG_FILE) if os.path.dirname(CONFIG_FILE) else ".", exist_ok=True)
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        if show_error and 'st' in globals():
            st.error(f"保存 API key 失败: {str(e)}")
        return False

def load_api_key() -> Optional[str]:
    """从本地配置文件加载 API key"""
    try:
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)
                api_key = decode_api_key(config.get("api_key", ""))
                if api_key:
                    return api_key
    except Exception as e:
        # 静默失败，如果文件损坏或不存在，返回 None
        pass
    return None

def load_embedding_model_config() -> str:
    """从本地配置文件加载嵌入模型配置
    
    Returns:
        模型名称，默认为 "BAAI/bge-small-zh-v1.5"
    """
    try:
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)
                model_name = config.get("embedding_model", "BAAI/bge-small-zh-v1.5")
                return model_name
    except Exception:
        pass
    return "BAAI/bge-small-zh-v1.5"

def save_embedding_model_config(model_name: str) -> bool:
    """保存嵌入模型配置到本地配置文件
    
    Args:
        model_name: 模型名称
    
    Returns:
        是否保存成功
    """
    try:
        # 读取现有配置
        config = {}
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
            except:
                pass
        
        # 更新嵌入模型配置
        config["embedding_model"] = model_name
        config["embedding_model_updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 保存配置
        os.makedirs(os.path.dirname(CONFIG_FILE) if os.path.dirname(CONFIG_FILE) else ".", exist_ok=True)
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        if 'st' in globals():
            st.error(f"保存嵌入模型配置失败: {str(e)}")
        return False

def download_model(model_name: str, progress_callback=None) -> bool:
    """下载HuggingFace模型
    
    Args:
        model_name: 模型名称（如 "BAAI/bge-base-zh-v1.5"）
        progress_callback: 进度回调函数，接收 (progress, message) 参数
    
    Returns:
        是否下载成功
    """
    try:
        try:
            from huggingface_hub import snapshot_download
        except ImportError:
            # 如果没有huggingface_hub，尝试使用transformers
            try:
                from transformers import AutoModel, AutoTokenizer
                if progress_callback:
                    progress_callback(50, f"🔄 正在下载模型 {model_name}...")
                # 使用transformers下载（会自动缓存）
                AutoModel.from_pretrained(model_name)
                AutoTokenizer.from_pretrained(model_name)
                if progress_callback:
                    progress_callback(100, f"✅ 模型 {model_name} 下载完成！")
                return True
            except ImportError:
                if progress_callback:
                    progress_callback(100, f"❌ 请安装 huggingface_hub 或 transformers: pip install huggingface_hub")
                return False
        
        if progress_callback:
            progress_callback(10, f"🔄 开始下载模型 {model_name}...")
        
        # 使用huggingface_hub下载
        cache_dir = os.path.join(
            os.path.expanduser("~"),
            ".cache",
            "huggingface",
            "hub"
        )
        
        if progress_callback:
            progress_callback(30, f"🔄 正在下载模型文件（这可能需要几分钟）...")
        
        snapshot_download(
            repo_id=model_name,
            cache_dir=cache_dir,
            local_files_only=False
        )
        
        if progress_callback:
            progress_callback(100, f"✅ 模型 {model_name} 下载完成！")
        
        return True
    except Exception as e:
        if progress_callback:
            progress_callback(100, f"❌ 下载失败: {str(e)}")
        return False

def delete_api_key(show_error: bool = True) -> bool:
    """删除本地保存的 API key"""
    try:
        if os.path.exists(CONFIG_FILE):
            os.remove(CONFIG_FILE)
            return True
        return False
    except Exception as e:
        if show_error and 'st' in globals():
            st.error(f"删除 API key 失败: {str(e)}")
        return False

# 文件读取模块
def read_text_file(file_path):
    """读取txt文件"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def read_docx_file(file_path):
    """读取Word文档（包括段落和表格）"""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():  # 只添加非空段落
                text_parts.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(' | '.join(row_text))
            if table_text:
                text_parts.append('\n表格:\n' + '\n'.join(table_text))
        
        return '\n\n'.join(text_parts) if text_parts else ""
    except ImportError:
        return "请安装python-docx: pip install python-docx"
    except Exception as e:
        return f"Word文档读取失败: {str(e)}"

def read_pdf_file(file_path):
    """读取PDF文件"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"PDF读取失败: {str(e)}"

def read_excel_file(file_path):
    """读取Excel文件"""
    try:
        import pandas as pd
        excel_content = {}
        xls = pd.ExcelFile(file_path)
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            excel_content[sheet_name] = df.to_string()
        return excel_content
    except ImportError:
        return {"错误": "请安装pandas和openpyxl"}

def read_markdown_file(file_path):
    """读取Markdown文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"Markdown读取失败: {str(e)}"

def read_javascript_file(file_path):
    """读取JavaScript文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"JavaScript读取失败: {str(e)}"

def read_json_file(file_path):
    """读取JSON文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = json.load(f)
            # 将JSON格式化为易读的字符串
            return json.dumps(content, ensure_ascii=False, indent=2)
    except json.JSONDecodeError as e:
        return f"JSON解析失败: {str(e)}"
    except Exception as e:
        return f"JSON读取失败: {str(e)}"

def process_folder(folder_path: str) -> Dict[str, Any]:
    """处理文件夹中的所有文件"""
    all_docs = {}
    
    # 支持的文件类型
    file_patterns = {
        '*.txt': ('txt', read_text_file),
        '*.docx': ('docx', read_docx_file),
        '*.pdf': ('pdf', read_pdf_file),
        '*.xlsx': ('excel', read_excel_file),
        '*.xls': ('excel', read_excel_file),
        '*.md': ('markdown', read_markdown_file),
        '*.js': ('javascript', read_javascript_file),
        '*.json': ('json', read_json_file),
    }
    
    for pattern, (file_type, reader_func) in file_patterns.items():
        for file_path in glob.glob(os.path.join(folder_path, pattern)):
            file_name = os.path.basename(file_path)
            
            # 跳过临时文件和隐藏文件
            # Excel 临时文件以 ~$ 开头，Word 临时文件也可能以 ~$ 开头
            if file_name.startswith('~$') or file_name.startswith('.'):
                continue
            
            try:
                content = reader_func(file_path)
                all_docs[file_name] = {
                    'path': file_path,
                    'content': content,
                    'type': file_type,
                    'size': os.path.getsize(file_path)
                }
            except Exception as e:
                all_docs[file_name] = {
                    'path': file_path,
                    'content': f"读取失败: {str(e)}",
                    'type': 'error',
                    'size': 0
                }
    
    return all_docs

# 本地向量数据库模块（不需要API密钥）
def get_model_path(model_name: str = "BAAI/bge-small-zh-v1.5") -> str:
    """获取模型路径，优先使用本地路径
    
    Args:
        model_name: HuggingFace 模型名称或本地路径
    
    Returns:
        模型路径（本地路径如果存在，否则返回模型名称）
    """
    import os
    
    try:
        # 如果已经是本地路径且存在，直接返回
        if os.path.exists(model_name) and os.path.isdir(model_name):
            return model_name
        
        # 检查 HuggingFace 缓存目录
        cache_dir = os.path.join(
            os.path.expanduser("~"),
            ".cache",
            "huggingface",
            "hub"
        )
        
        # 将模型名称转换为缓存目录格式（BAAI/bge-small-zh-v1.5 -> models--BAAI--bge-small-zh-v1.5）
        cache_model_name = f"models--{model_name.replace('/', '--')}"
        cache_path = os.path.join(cache_dir, cache_model_name)
        
        # 查找 snapshots 目录下的最新版本
        if os.path.exists(cache_path):
            snapshots_dir = os.path.join(cache_path, "snapshots")
            if os.path.exists(snapshots_dir):
                try:
                    # 先尝试快速检查目录是否可访问，避免卡住
                    if not os.access(snapshots_dir, os.R_OK):
                        raise PermissionError("No read permission")
                    
                    # 使用 Path 对象，通常比 os.listdir 更安全，避免在大型目录上卡住
                    from pathlib import Path
                    snapshots_path = Path(snapshots_dir)
                    # 使用 try-except 包裹 iterdir()，避免在某些文件系统上卡住
                    try:
                        snapshots = [d.name for d in snapshots_path.iterdir() 
                                   if d.is_dir()]
                    except (OSError, PermissionError):
                        # 如果 iterdir() 失败，回退到 os.listdir，但限制数量
                        try:
                            all_items = os.listdir(snapshots_dir)
                            snapshots = [d for d in all_items 
                                       if os.path.isdir(os.path.join(snapshots_dir, d))]
                        except (OSError, PermissionError):
                            snapshots = []
                    
                    if snapshots:
                        # 使用最新的快照
                        latest_snapshot = sorted(snapshots)[-1]
                        local_path = os.path.join(snapshots_dir, latest_snapshot)
                        if os.path.exists(local_path):
                            return local_path
                except (OSError, PermissionError, Exception):
                    pass  # 忽略所有错误，继续检查其他路径
        
        # 检查项目目录下的 models 文件夹
        project_model_path = os.path.join(".", "models", model_name.replace("/", "--"))
        if os.path.exists(project_model_path):
            return project_model_path
    except Exception:
        pass  # 如果出现任何错误，返回原始模型名称
    
    # 如果都不存在，返回原始模型名称（会触发下载）
    return model_name

def check_db_corrupted(db_path: str) -> bool:
    """检测向量数据库是否损坏（特别是 schema 兼容性问题）
    
    Args:
        db_path: 向量数据库路径
    
    Returns:
        True 如果数据库损坏，False 如果正常或不存在
    """
    if not os.path.exists(db_path):
        return False
    
    try:
        # 尝试加载数据库来检测是否损坏
        try:
            from langchain_huggingface import HuggingFaceEmbeddings
        except ImportError:
            from langchain_community.embeddings import HuggingFaceEmbeddings
        
        try:
            from langchain_chroma import Chroma
        except ImportError:
            from langchain_community.vectorstores import Chroma
        
        # 初始化嵌入模型
        embedding_model = load_embedding_model_config()
        model_path = get_model_path(embedding_model)
        embeddings = HuggingFaceEmbeddings(
            model_name=model_path,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        
        # 尝试加载向量数据库
        vectorstore = Chroma(
            persist_directory=db_path,
            embedding_function=embeddings
        )
        
        # 尝试访问数据库（如果 schema 错误会在这里失败）
        _ = len(vectorstore)
        return False  # 数据库正常
    except Exception as e:
        error_msg = str(e).lower()
        # 只检测真正表示数据库损坏的关键错误信息
        # 避免过于宽泛的匹配，防止误判正常错误
        is_corrupted = (
            # Schema 相关错误（版本兼容性问题）
            "no such column" in error_msg or
            "collections.topic" in error_msg or
            ("schema" in error_msg and ("mismatch" in error_msg or "invalid" in error_msg or "version" in error_msg)) or
            # HNSW 索引损坏（明确的错误信息）
            ("hnsw" in error_msg and ("corrupt" in error_msg or "invalid" in error_msg or "damaged" in error_msg)) or
            # 索引文件损坏（明确的错误信息）
            ("index" in error_msg and ("corrupt" in error_msg or "invalid" in error_msg or "damaged" in error_msg or "missing" in error_msg)) or
            # SQLite 数据库文件损坏（明确的错误信息）
            ("sqlite" in error_msg and ("corrupt" in error_msg or "database disk image is malformed" in error_msg or "file is encrypted" in error_msg)) or
            # 段文件损坏
            ("segment" in error_msg and ("corrupt" in error_msg or "invalid" in error_msg or "damaged" in error_msg))
        )
        if is_corrupted:
            print(f"⚠️ 检测到数据库损坏: {str(e)}")
        else:
            # 记录非损坏性错误（用于调试）
            print(f"ℹ️ 数据库加载时出现非损坏性错误（将忽略）: {str(e)}")
        return is_corrupted

def cleanup_corrupted_db(db_path: str, force: bool = True):
    """彻底清理损坏的向量数据库目录
    
    Args:
        db_path: 向量数据库路径
        force: 是否强制清理（包括多次尝试和延迟）
    
    Returns:
        bool: 是否成功清理
    """
    import shutil
    import time
    
    if not os.path.exists(db_path):
        return True
    
    if force:
        # 强制清理模式：多次尝试，确保彻底删除（Windows上需要更长时间）
        import platform
        is_windows = platform.system() == 'Windows'
        max_attempts = 8 if is_windows else 5  # Windows上增加重试次数
        wait_time = 2.0 if is_windows else 0.5  # Windows上增加等待时间
        
        for attempt in range(max_attempts):
            try:
                # 先尝试正常删除
                if os.path.exists(db_path):
                    shutil.rmtree(db_path, ignore_errors=False)
                
                # 等待一下，确保文件系统更新（Windows需要更长时间）
                time.sleep(wait_time)
                
                # 验证是否删除成功
                if not os.path.exists(db_path):
                    print(f"[OK] 已彻底清理损坏的向量数据库目录: {db_path}")
                    return True
                    
            except PermissionError as pe:
                # Windows 上可能有文件被锁定，等待后重试
                if attempt < max_attempts - 1:
                    wait_interval = wait_time * (attempt + 1)  # 递增等待时间
                    print(f"[WARN] 文件被锁定，等待 {wait_interval:.1f} 秒后重试 ({attempt + 1}/{max_attempts})...")
                    time.sleep(wait_interval)
                    continue
                else:
                    print(f"[ERROR] 清理失败（文件被锁定）: {str(pe)}")
                    print(f"   请手动删除目录: {db_path}")
                    return False
            except Exception as e:
                if attempt < max_attempts - 1:
                    wait_interval = wait_time * (attempt + 1)
                    print(f"[WARN] 清理失败，等待 {wait_interval:.1f} 秒后重试 ({attempt + 1}/{max_attempts}): {str(e)}")
                    time.sleep(wait_interval)
                    continue
                else:
                    print(f"[ERROR] 清理失败: {str(e)}")
                    print(f"   请手动删除目录: {db_path}")
                    return False
        
        # 如果多次尝试后仍然存在，尝试使用 ignore_errors（忽略错误强制删除）
        if os.path.exists(db_path):
            try:
                print(f"[INFO] 尝试强制删除模式（忽略错误）...")
                shutil.rmtree(db_path, ignore_errors=True)
                time.sleep(wait_time * 2)  # 等待更长时间
                if not os.path.exists(db_path):
                    print(f"[OK] 已强制清理数据库目录: {db_path}")
                    return True
            except Exception as e:
                print(f"[WARN] 强制删除模式也失败: {str(e)}")
        
        # 最后尝试：重命名目录（如果无法删除，至少不影响后续操作）
        if os.path.exists(db_path):
            try:
                temp_name = db_path + "_deleted_" + str(int(time.time()))
                os.rename(db_path, temp_name)
                print(f"[WARN] 无法删除目录，已重命名为: {temp_name}")
                print(f"   可以在程序关闭后手动删除该备份目录")
                return True  # 重命名成功也算成功（不影响后续操作）
            except Exception as e:
                print(f"❌ 无法清理或重命名目录: {str(e)}")
                print(f"   目录可能被其他进程占用，请手动删除: {db_path}")
                return False
        
        return False
    else:
        # 简单清理模式
        try:
            if os.path.exists(db_path):
                shutil.rmtree(db_path)
                print(f"[OK] 已清理数据库目录: {db_path}")
                return True
        except Exception as e:
            print(f"⚠️ 清理失败: {str(e)}")
            return False
    
    return False

def get_vector_db_path(folder_path: str) -> str:
    """根据文件夹路径生成唯一的向量数据库目录路径
    
    Args:
        folder_path: 文件夹路径
    
    Returns:
        向量数据库目录路径
    """
    if not folder_path:
        # 上传文件时没有文件夹路径，使用默认路径
        return "./chroma_db"
    
    # 使用路径的哈希值创建唯一目录名
    # 规范化路径：与 normalize_path 保持一致，确保路径规范化逻辑统一
    try:
        # 先规范化路径格式（统一斜杠）
        normalized = os.path.normpath(folder_path)
        # 转换为绝对路径（如果路径存在）
        if os.path.exists(normalized):
            normalized = os.path.abspath(normalized)
        else:
            # 如果路径不存在，仍然规范化格式
            normalized = os.path.normpath(normalized)
        # 统一使用正斜杠，并转换为小写（Windows路径大小写不敏感）
        normalized_path = normalized.replace('\\', '/')
        if os.name == 'nt':  # Windows系统
            normalized_path = normalized_path.lower()
    except Exception as e:
        # 如果规范化失败，使用基本规范化
        print(f"⚠️ 路径规范化失败: {folder_path}, 错误: {str(e)}")
        normalized_path = os.path.normpath(folder_path).replace('\\', '/')
        if os.name == 'nt':
            normalized_path = normalized_path.lower()
    
    # 使用规范化后的路径计算哈希值
    path_hash = hashlib.md5(normalized_path.encode('utf-8')).hexdigest()[:12]
    
    # 创建安全的目录名（移除特殊字符）
    safe_folder_name = os.path.basename(normalized_path)
    safe_folder_name = "".join(c for c in safe_folder_name if c.isalnum() or c in (' ', '-', '_'))[:30]
    safe_folder_name = safe_folder_name.strip() or "unknown"
    
    # 组合目录名：文件夹名_哈希值
    db_dir_name = f"{safe_folder_name}_{path_hash}"
    return os.path.join("./chroma_db", db_dir_name)

def load_existing_vector_store(folder_path: str = None, progress_callback=None):
    """加载已有的向量数据库
    
    Args:
        folder_path: 文件夹路径（用于确定向量数据库位置）
        progress_callback: 进度回调函数，接收 (progress, message) 参数
    
    Returns:
        向量数据库对象，如果不存在或加载失败则返回 None
    """
    try:
        # 优先使用新版本的包
        try:
            from langchain_huggingface import HuggingFaceEmbeddings
        except ImportError:
            from langchain_community.embeddings import HuggingFaceEmbeddings
        
        try:
            from langchain_chroma import Chroma
        except ImportError:
            from langchain_community.vectorstores import Chroma
        
        db_path = get_vector_db_path(folder_path) if folder_path else "./chroma_db"
        
        if not os.path.exists(db_path):
            return None
        
        if progress_callback:
            progress_callback(10, "🔄 正在加载已有向量数据库...")
        
        # 初始化嵌入模型（必须与创建时使用相同的模型）
        # 优先使用本地模型路径，避免网络下载
        embedding_model = load_embedding_model_config()
        model_path = get_model_path(embedding_model)
        embeddings = HuggingFaceEmbeddings(
            model_name=model_path,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        
        if progress_callback:
            progress_callback(50, "🔄 正在加载向量数据库...")
        
        # 从持久化目录加载
        vectorstore = Chroma(
            persist_directory=db_path,
            embedding_function=embeddings
        )
        
        # 验证向量数据库是否可用（使用更温和的验证方法）
        # 只对真正的schema错误或维度不匹配错误进行清理，其他错误只记录但不清理
        try:
            # 尝试获取数据库中的文档数量
            doc_count = len(vectorstore)
            
            # 如果文档数量为0，可能是空数据库，但不算损坏
            if doc_count == 0:
                if progress_callback:
                    progress_callback(100, "⚠️ 向量数据库为空，将重新创建...")
                return None
            
            # 尝试进行一次简单的查询来验证数据库是否真的可用
            # 使用一个简单的测试查询，如果失败说明数据库有问题
            try:
                # 尝试获取第一个文档（如果存在）
                results = vectorstore.similarity_search("test", k=1)
                # 如果能正常返回结果（即使为空），说明数据库可用
            except Exception as query_error:
                # 查询失败，检查是否是schema错误
                error_msg = str(query_error).lower()
                is_schema_error = (
                    "no such column" in error_msg or
                    "collections.topic" in error_msg or
                    "schema" in error_msg or
                    "dimensionality" in error_msg or
                    "dimension" in error_msg
                )
                
                if is_schema_error:
                    # Schema错误或维度不匹配，清理数据库
                    if progress_callback:
                        progress_callback(100, "⚠️ 检测到数据库 schema 错误或维度不匹配，正在清理...")
                    cleanup_corrupted_db(db_path, force=True)
                    return None
                else:
                    # 其他查询错误，可能是临时性问题，不清理数据库，但返回None让调用者重新创建
                    # 记录错误但不清理，因为可能是临时性问题
                    print(f"⚠️ 向量数据库查询失败（可能是临时性问题）: {str(query_error)}")
                    if progress_callback:
                        progress_callback(100, "⚠️ 向量数据库查询失败，将重新创建...")
                    return None
                    
        except Exception as verify_error:
            # len() 调用失败，检查是否是schema错误
            error_msg = str(verify_error).lower()
            is_schema_error = (
                "no such column" in error_msg or
                "collections.topic" in error_msg or
                "schema" in error_msg or
                "dimensionality" in error_msg or
                "dimension" in error_msg
            )
            
            if is_schema_error:
                # Schema错误或维度不匹配，清理数据库
                if progress_callback:
                    progress_callback(100, "⚠️ 检测到数据库 schema 错误或维度不匹配，正在清理...")
                cleanup_corrupted_db(db_path, force=True)
            else:
                # 其他错误（可能是临时性问题），不清理数据库，只返回None
                # 记录错误但不清理，因为可能是临时性问题
                print(f"⚠️ 向量数据库验证失败（可能是临时性问题）: {str(verify_error)}")
                if progress_callback:
                    progress_callback(100, "⚠️ 向量数据库验证失败，将重新创建...")
            return None
        
        if progress_callback:
            progress_callback(100, "✅ 向量数据库加载完成！")
        
        return vectorstore
    except Exception as e:
        # 加载失败，返回 None
        return None

def calculate_content_hash(content: Any) -> str:
    """计算文档内容的哈希值
    
    Args:
        content: 文档内容（字符串或字典）
    
    Returns:
        内容的 MD5 哈希值
    """
    if isinstance(content, dict):
        # Excel 文件：合并所有工作表内容
        content_str = "\n".join([f"{k}:{v}" for k, v in content.items()])
    else:
        content_str = str(content)
    
    return hashlib.md5(content_str.encode('utf-8')).hexdigest()

def normalize_path(path: str) -> str:
    """规范化路径，用于比较"""
    if not path:
        return ""
    # 统一使用 os.path.normpath 和 os.path.abspath 来规范化路径
    try:
        # 先规范化路径格式（统一斜杠）
        normalized = os.path.normpath(path)
        # 转换为绝对路径（如果可能）
        if os.path.exists(normalized):
            normalized = os.path.abspath(normalized)
        else:
            # 如果路径不存在，仍然规范化格式
            normalized = os.path.normpath(normalized)
        # 统一转换为小写（Windows路径大小写不敏感）
        if os.name == 'nt':  # Windows
            normalized = normalized.lower()
        return normalized
    except Exception as e:
        # 如果规范化失败，至少统一格式
        print(f"⚠️ 路径规范化失败: {path}, 错误: {str(e)}")
        return os.path.normpath(path).lower() if os.name == 'nt' else os.path.normpath(path)

def check_docs_changed(docs_dict: Dict[str, Any], folder_path: str) -> bool:
    """检查文档是否发生变化（包括模型变化）
    
    Args:
        docs_dict: 当前文档字典
        folder_path: 文件夹路径
    
    Returns:
        True 如果文档或模型发生变化，False 如果未变化
    """
    # 创建文档签名文件路径（基于文件夹路径）
    db_path = get_vector_db_path(folder_path)
    signature_file = os.path.join(db_path, ".docs_signature.json")
    
    # 检查数据库目录是否存在
    if not os.path.exists(db_path):
        print(f"[INFO] 向量数据库目录不存在: {db_path}")
        return True  # 数据库目录不存在，认为需要创建
    
    # 检查签名文件是否存在
    if not os.path.exists(signature_file):
        print(f"[INFO] 文档签名文件不存在: {signature_file}")
        return True  # 签名文件不存在，认为文档已变化
    
    try:
        # 读取之前的签名
        with open(signature_file, 'r', encoding='utf-8') as f:
            old_signature = json.load(f)
        
        # 检查嵌入模型是否变化
        current_embedding_model = load_embedding_model_config()
        old_embedding_model = old_signature.get("embedding_model", "BAAI/bge-small-zh-v1.5")
        old_embedding_dimension = old_signature.get("embedding_dimension", 384)
        current_embedding_dimension = get_embedding_model_dimension(current_embedding_model)
        
        if old_embedding_model != current_embedding_model or old_embedding_dimension != current_embedding_dimension:
            print(f"[CHANGE] 嵌入模型变化: {old_embedding_model} ({old_embedding_dimension}维) -> {current_embedding_model} ({current_embedding_dimension}维)")
            return True
        
        # 生成当前文档签名
        # 使用规范化后的路径，确保与保存的签名路径格式一致
        normalized_current_folder_path = normalize_path(folder_path) if folder_path else None
        
        current_signature = {
            "folder_path": normalized_current_folder_path,  # 使用规范化后的路径
            "file_count": len(docs_dict),
            "files": {}
        }
        
        for filename, data in docs_dict.items():
            file_path = data.get('path', '')
            
            file_info = {
                "size": data.get('size', 0),
                "type": data.get('type', '')  # 文件类型
            }
            
            # 如果文件路径存在且是持久路径（非临时路径），记录完整路径和修改时间
            if file_path and os.path.exists(file_path):
                # 检查是否是临时路径（临时路径通常包含 temp 或 tmp）
                is_temp_path = 'temp' in file_path.lower() or 'tmp' in file_path.lower()
                if not is_temp_path:
                    # 保存完整路径（规范化后）和修改时间
                    file_info["path"] = normalize_path(file_path)  # 保存规范化后的完整路径
                    file_info["mtime"] = os.path.getmtime(file_path)
            
            current_signature["files"][filename] = file_info
        
        # 比较签名 - 规范化路径后再比较
        old_folder_path = old_signature.get("folder_path")
        current_folder_path = current_signature["folder_path"]
        
        # 规范化路径进行比较（处理 None、空字符串、路径格式差异等情况）
        old_path_norm = normalize_path(old_folder_path) if old_folder_path else ""
        current_path_norm = normalize_path(current_folder_path) if current_folder_path else ""
        
        # 路径比较逻辑：
        # 1. 如果两个路径都不为空且不同，才认为路径变化
        # 2. 如果旧路径为空但新路径不为空，可能是首次创建，继续检查文件内容
        # 3. 如果旧路径不为空但新路径为空，可能是路径丢失，继续检查文件内容（如果文件内容匹配，仍可使用）
        # 4. 如果两个路径都为空，继续检查文件内容
        if old_path_norm and current_path_norm:
            # 两个路径都不为空，需要比较
            if old_path_norm != current_path_norm:
                print(f"[CHANGE] 文件夹路径变化: {old_folder_path} -> {current_folder_path}")
                return True
            # 路径相同，继续检查文件内容
        elif not old_path_norm and current_path_norm:
            # 旧路径为空但新路径不为空，可能是首次创建
            # 如果文件内容哈希都匹配，可以认为未变化（兼容旧版本）
            print(f"[INFO] 文件夹路径状态变化: 旧路径为空, 新路径={current_folder_path}")
            # 继续检查文件内容
        elif old_path_norm and not current_path_norm:
            # 旧路径不为空但新路径为空，可能是路径丢失
            # 如果文件内容哈希都匹配，仍可使用（兼容性）
            print(f"[INFO] 文件夹路径状态变化: 旧路径={old_folder_path}, 新路径为空")
            # 继续检查文件内容
        # 如果两个路径都为空，继续检查文件内容
        
        if old_signature.get("file_count") != current_signature["file_count"]:
            print(f"[CHANGE] 文件数量变化: {old_signature.get('file_count')} -> {current_signature['file_count']}")
            return True
        
        old_files = old_signature.get("files", {})
        current_files = current_signature["files"]
        
        if set(old_files.keys()) != set(current_files.keys()):
            old_keys = set(old_files.keys())
            current_keys = set(current_files.keys())
            added = current_keys - old_keys
            removed = old_keys - current_keys
            print(f"[CHANGE] 文件名变化: 新增={added}, 删除={removed}")
            return True
        
        # 检查文件：文件名（完整路径）、大小、类型、修改时间
        changed_files = []
        for filename in old_files.keys():
            if filename not in current_files:
                changed_files.append(f"{filename} (已删除)")
                return True
            
            old_info = old_files[filename]
            current_info = current_files[filename]
            
            # 1. 检查文件大小
            if old_info.get("size") != current_info.get("size"):
                changed_files.append(f"{filename} (大小变化: {old_info.get('size')} -> {current_info.get('size')})")
                return True
            
            # 2. 检查文件类型
            if old_info.get("type") != current_info.get("type"):
                changed_files.append(f"{filename} (类型变化: {old_info.get('type')} -> {current_info.get('type')})")
                return True
            
            # 3. 检查文件路径（完整路径）
            old_path = old_info.get("path")
            current_path = current_info.get("path")
            if old_path and current_path:
                # 两者都有路径，比较规范化后的路径
                if old_path != current_path:
                    changed_files.append(f"{filename} (路径变化: {old_path} -> {current_path})")
                    return True
            
            # 4. 检查修改时间（如果都存在）
            old_mtime = old_info.get("mtime")
            current_mtime = current_info.get("mtime")
            if old_mtime and current_mtime:
                # 修改时间差异超过1秒认为文件已变化
                if abs(old_mtime - current_mtime) > 1:
                    changed_files.append(f"{filename} (修改时间变化)")
                    return True
        
        if changed_files:
            print(f"[CHANGE] 检测到文档变化: {', '.join(changed_files)}")
            return True
        
        # 文档未变化，只检查数据库文件是否存在，不验证是否可用
        # 如果文档未变化，即使数据库损坏，也不应该自动重新创建
        # 数据库的可用性由 load_existing_vector_store 来验证
        try:
            # 只检查数据库文件是否存在
            chroma_sqlite = os.path.join(db_path, "chroma.sqlite3")
            if not os.path.exists(chroma_sqlite):
                print(f"[INFO] 数据库文件不存在: {chroma_sqlite}")
                return True  # 数据库文件不存在，需要重新创建
            
            # 检查数据库文件大小，如果为0，可能损坏
            if os.path.getsize(chroma_sqlite) == 0:
                print(f"[INFO] 数据库文件为空，可能损坏")
                return True  # 数据库文件为空，需要重新创建
            
            # 文档未变化且数据库文件存在，返回 False（文档未变化）
            # 数据库的可用性由 load_existing_vector_store 来验证
            # 如果数据库损坏，load_existing_vector_store 会返回 None，界面会提示但不会自动重新创建
            print("[OK] 文档未变化，可以使用已有向量数据库")
            return False
        except Exception as e:
            print(f"[WARN] 验证数据库文件时出错: {str(e)}")
            # 验证出错，保守起见认为需要重新创建
            return True
    except Exception as e:
        # 读取签名失败，记录错误但认为文档已变化
        print(f"[ERROR] 读取文档签名失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return True

def get_embedding_model_dimension(model_name: str) -> int:
    """获取嵌入模型的向量维度
    
    Args:
        model_name: 模型名称
    
    Returns:
        向量维度
    """
    model_dimensions = {
        "BAAI/bge-small-zh-v1.5": 384,
        "BAAI/bge-base-zh-v1.5": 768,
        "BAAI/bge-large-zh-v1.5": 1024,
    }
    return model_dimensions.get(model_name, 384)  # 默认384

def save_docs_signature(docs_dict: Dict[str, Any], folder_path: str):
    """保存文档签名
    
    Args:
        docs_dict: 文档字典
        folder_path: 文件夹路径
    """
    try:
        db_path = get_vector_db_path(folder_path)
        os.makedirs(db_path, exist_ok=True)
        signature_file = os.path.join(db_path, ".docs_signature.json")
        
        # 获取当前使用的嵌入模型
        embedding_model = load_embedding_model_config()
        embedding_dimension = get_embedding_model_dimension(embedding_model)
        
        # 保存规范化后的路径，确保路径比较时一致
        normalized_folder_path = normalize_path(folder_path) if folder_path else None
        
        signature = {
            "folder_path": normalized_folder_path,  # 保存规范化后的路径
            "file_count": len(docs_dict),
            "files": {},
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "embedding_model": embedding_model,  # 保存使用的模型
            "embedding_dimension": embedding_dimension  # 保存模型维度
        }
        
        for filename, data in docs_dict.items():
            file_path = data.get('path', '')
            
            file_info = {
                "size": data.get('size', 0),
                "type": data.get('type', '')  # 文件类型
            }
            
            # 如果文件路径存在且是持久路径（非临时路径），记录完整路径和修改时间
            if file_path and os.path.exists(file_path):
                # 检查是否是临时路径
                is_temp_path = 'temp' in file_path.lower() or 'tmp' in file_path.lower()
                if not is_temp_path:
                    # 保存完整路径（规范化后）和修改时间
                    file_info["path"] = normalize_path(file_path)  # 保存规范化后的完整路径
                    file_info["mtime"] = os.path.getmtime(file_path)
            
            signature["files"][filename] = file_info
        
        with open(signature_file, 'w', encoding='utf-8') as f:
            json.dump(signature, f, indent=2, ensure_ascii=False)
    except Exception:
        pass  # 保存签名失败不影响主流程

class ProgressEmbeddings:
    """包装的嵌入模型类，用于在生成向量时更新进度"""
    def __init__(self, embeddings, progress_callback=None, total_docs=0, start_progress=70, end_progress=85):
        self.embeddings = embeddings
        self.progress_callback = progress_callback
        self.total_docs = total_docs
        self.start_progress = start_progress
        self.end_progress = end_progress
        self.processed_docs = 0
        self.batch_size = 50  # 内部批处理大小，用于进度更新
    
    def embed_documents(self, texts):
        """批量生成向量，并更新进度（如果一次性处理所有文档，则内部分批处理）"""
        total_texts = len(texts)
        
        # 如果文档数量较少，直接处理
        if total_texts <= self.batch_size:
            result = self.embeddings.embed_documents(texts)
            self.processed_docs += total_texts
            
            if self.progress_callback and self.total_docs > 0:
                progress = self.start_progress + int((self.processed_docs / self.total_docs) * (self.end_progress - self.start_progress))
                progress = min(progress, self.end_progress)
                self.progress_callback(progress, f"🔄 步骤 3/4: 生成向量嵌入 ({self.processed_docs}/{self.total_docs} 个文档块)...")
            
            return result
        
        # 如果文档数量较多，分批处理以显示进度
        all_embeddings = []
        num_batches = (total_texts + self.batch_size - 1) // self.batch_size
        
        for batch_idx in range(num_batches):
            start_idx = batch_idx * self.batch_size
            end_idx = min(start_idx + self.batch_size, total_texts)
            batch_texts = texts[start_idx:end_idx]
            
            # 生成当前批次的向量
            batch_embeddings = self.embeddings.embed_documents(batch_texts)
            all_embeddings.extend(batch_embeddings)
            
            # 更新进度
            self.processed_docs += len(batch_texts)
            if self.progress_callback and self.total_docs > 0:
                progress = self.start_progress + int((self.processed_docs / self.total_docs) * (self.end_progress - self.start_progress))
                progress = min(progress, self.end_progress)
                self.progress_callback(progress, f"🔄 步骤 3/4: 生成向量嵌入 ({self.processed_docs}/{self.total_docs} 个文档块)...")
        
        return all_embeddings
    
    def embed_query(self, text):
        """单个查询向量化（不更新进度）"""
        return self.embeddings.embed_query(text)
    
    def __getattr__(self, name):
        """代理其他属性和方法到原始 embeddings 对象"""
        return getattr(self.embeddings, name)

def create_local_vector_store(docs_dict: Dict[str, Any], progress_callback=None, folder_path: str = None):
    """创建本地向量数据库，使用开源嵌入模型
    
    Args:
        docs_dict: 文档字典
        progress_callback: 进度回调函数，接收 (progress, message) 参数
        folder_path: 文件夹路径（用于签名）
    """
    try:
        # 兼容不同版本的 langchain 导入
        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter
        except ImportError:
            try:
                from langchain.text_splitter import RecursiveCharacterTextSplitter
            except ImportError:
                from langchain_core.text_splitter import RecursiveCharacterTextSplitter
        
        # 优先使用新版本的包
        try:
            from langchain_huggingface import HuggingFaceEmbeddings
        except ImportError:
            from langchain_community.embeddings import HuggingFaceEmbeddings
        
        try:
            from langchain_chroma import Chroma
        except ImportError:
            from langchain_community.vectorstores import Chroma
        try:
            from langchain.schema import Document as LangDocument
        except ImportError:
            from langchain_core.documents import Document as LangDocument
        
        # 步骤 0: 在开始处理之前，先检测并清理损坏的数据库目录（避免版本兼容性问题）
        db_path = get_vector_db_path(folder_path) if folder_path else "./chroma_db"
        
        # 检查目录权限（在创建目录之前）
        parent_dir = os.path.dirname(db_path) if os.path.dirname(db_path) else "."
        if not os.path.exists(parent_dir):
            os.makedirs(parent_dir, exist_ok=True)
        if not os.access(parent_dir, os.W_OK):
            raise PermissionError(f"没有写入权限: {parent_dir}")
        
        # 如果数据库目录已存在，先检测是否损坏或模型维度不匹配
        # 注意：此函数只在文档变化或数据库不存在时被调用
        # 如果数据库存在且正常，调用者应该已经检查过文档变化
        if os.path.exists(db_path):
            if progress_callback:
                progress_callback(5, "🔄 检测向量数据库状态...")
            
            # 检查模型维度是否匹配
            embedding_model = load_embedding_model_config()
            expected_dimension = get_embedding_model_dimension(embedding_model)
            
            # 尝试检测现有数据库的维度
            dimension_mismatch = False
            try:
                # 尝试加载数据库来检测维度
                try:
                    from langchain_huggingface import HuggingFaceEmbeddings
                except ImportError:
                    from langchain_community.embeddings import HuggingFaceEmbeddings
                
                try:
                    from langchain_chroma import Chroma
                except ImportError:
                    from langchain_community.vectorstores import Chroma
                
                # 使用当前模型初始化
                model_path = get_model_path(embedding_model)
                test_embeddings = HuggingFaceEmbeddings(
                    model_name=model_path,
                    model_kwargs={'device': 'cpu'},
                    encode_kwargs={'normalize_embeddings': True}
                )
                
                # 尝试加载向量数据库
                test_vectorstore = Chroma(
                    persist_directory=db_path,
                    embedding_function=test_embeddings
                )
                
                # 尝试访问数据库，如果维度不匹配会抛出异常
                try:
                    _ = len(test_vectorstore)
                except Exception as dim_error:
                    error_msg = str(dim_error).lower()
                    if "dimension" in error_msg or "dimensionality" in error_msg:
                        dimension_mismatch = True
            except Exception:
                # 如果检测失败，假设维度不匹配（安全起见）
                dimension_mismatch = True
            
            # 检测数据库是否损坏（特别是 schema 兼容性问题）
            is_corrupted = check_db_corrupted(db_path)
            
            if is_corrupted or dimension_mismatch:
                # 数据库损坏或维度不匹配，需要清理后重新创建
                if dimension_mismatch:
                    if progress_callback:
                        progress_callback(5, f"⚠️ 检测到模型维度不匹配（当前模型维度: {expected_dimension}），正在清理旧数据库...")
                else:
                    if progress_callback:
                        progress_callback(5, "⚠️ 检测到数据库损坏（可能是版本兼容性问题），正在清理...")
                cleanup_corrupted_db(db_path, force=True)
                import time
                time.sleep(2)  # 增加等待时间，确保文件系统完全更新
                # 再次确认目录已删除
                if os.path.exists(db_path):
                    import shutil
                    try:
                        shutil.rmtree(db_path)
                        print(f"✅ 强制删除数据库目录: {db_path}")
                    except Exception as cleanup_error:
                        print(f"⚠️ 强制删除失败: {str(cleanup_error)}")
                        # 如果删除失败，重命名目录
                        backup_name = db_path + "_backup_" + str(int(time.time()))
                        try:
                            os.rename(db_path, backup_name)
                            print(f"⚠️ 已重命名数据库目录为备份: {backup_name}")
                        except:
                            pass
            else:
                # 数据库正常，但由于文档变化需要重新创建，清理旧数据库
                # 注意：调用者应该已经检查过文档变化，这里直接清理即可
                if progress_callback:
                    progress_callback(5, "📝 检测到文档变化，清理旧向量数据库...")
                cleanup_corrupted_db(db_path, force=True)
                import time
                time.sleep(0.5)  # 等待文件系统更新
        
        # 提取文本内容
        if progress_callback:
            progress_callback(15, "🔄 步骤 1/4: 提取文档内容...")
        
        texts = []
        total_files = len(docs_dict)
        for idx, (filename, data) in enumerate(docs_dict.items()):
            content = data['content']
            if isinstance(content, dict):  # Excel文件
                for sheet, sheet_content in content.items():
                    texts.append(f"文件: {filename} | 工作表: {sheet}\n{sheet_content}")
            else:
                texts.append(f"文件: {filename}\n{content}")
            
            if progress_callback and total_files > 0:
                progress = 15 + int((idx + 1) / total_files * 10)
                progress_callback(progress, f"🔄 步骤 1/4: 提取文档内容... ({idx + 1}/{total_files})")
        
        # 分割文本
        if progress_callback:
            progress_callback(30, "🔄 步骤 2/4: 分割文本...")
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        documents = []
        total_texts = len(texts)
        for i, text in enumerate(texts):
            splits = text_splitter.split_text(text)
            for split in splits:
                documents.append(LangDocument(
                    page_content=split,
                    metadata={"source": list(docs_dict.keys())[i % len(docs_dict)]}
                ))
            
            if progress_callback and total_texts > 0:
                progress = 30 + int((i + 1) / total_texts * 20)
                progress_callback(progress, f"🔄 步骤 2/4: 分割文本... ({i + 1}/{total_texts})")
        
        # 使用本地嵌入模型
        if progress_callback:
            progress_callback(55, "🔄 步骤 3/4: 初始化嵌入模型（首次运行会下载模型，可能需要几分钟）...")
        
        # 优先使用本地模型路径，避免网络下载
        embedding_model = load_embedding_model_config()
        model_path = get_model_path(embedding_model)
        
        try:
            embeddings = HuggingFaceEmbeddings(
                model_name=model_path,  # 使用配置的嵌入模型
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
        except Exception as model_error:
            error_type = type(model_error).__name__
            error_msg = str(model_error)
            raise Exception(f"嵌入模型初始化失败 [{error_type}]: {error_msg}\n\n"
                          f"可能的原因:\n"
                          f"- 模型文件不存在或损坏\n"
                          f"- 网络连接问题（无法下载模型）\n"
                          f"- 模型路径配置错误: {model_path}\n\n"
                          f"解决方案:\n"
                          f"- 检查模型文件是否存在\n"
                          f"- 使用 download_model.py 手动下载模型\n"
                          f"- 检查网络连接") from model_error
        
        # 检查文档是否为空
        if not documents or len(documents) == 0:
            raise ValueError("没有可用的文档内容，无法创建向量数据库。请检查文档是否为空或格式是否正确。")
        
        # 确保目录不存在后再创建（之前的清理应该已经删除了目录）
        # 如果目录仍然存在，强制删除（可能是维度不匹配的旧数据库）
        if os.path.exists(db_path):
            import shutil
            import time
            
            # 使用cleanup_corrupted_db函数进行强力清理
            cleanup_success = cleanup_corrupted_db(db_path, force=True)
            
            if not cleanup_success:
                # 如果清理失败，尝试重命名作为备份（避免阻塞后续操作）
                backup_name = db_path + "_backup_" + str(int(time.time()))
                try:
                    os.rename(db_path, backup_name)
                    print(f"⚠️ 无法删除目录，已重命名为备份: {backup_name}")
                    print(f"   可以在程序关闭后手动删除该备份目录")
                except Exception as rename_error:
                    # 如果重命名也失败，记录警告但不抛出异常（允许继续创建新数据库）
                    print(f"⚠️ 无法删除或重命名目录: {str(rename_error)}")
                    print(f"   目录 {db_path} 可能被其他进程占用")
                    print(f"   建议：关闭其他可能使用该数据库的程序，然后手动删除目录")
                    # 尝试创建一个带时间戳的新目录名
                    db_path = db_path + "_new_" + str(int(time.time()))
                    print(f"   将使用新目录: {db_path}")
        
        # 创建新目录
        os.makedirs(db_path, exist_ok=True)
        
        # 步骤 3: 预生成向量嵌入（带进度更新，用于显示进度和预热模型）
        # 注意：Chroma.from_documents 会在内部重新生成向量，但预生成可以：
        # 1. 显示详细的进度更新
        # 2. 预热模型（首次使用时会加载模型到内存）
        # 3. 提前发现模型错误
        total_docs = len(documents)
        batch_size = 50  # 每批处理50个文档，避免内存占用过大
        
        # 步骤 3: 使用带进度更新的嵌入模型
        total_docs = len(documents)
        if progress_callback:
            progress_callback(70, f"🔄 步骤 3/4: 生成向量嵌入（共 {total_docs} 个文档块，这可能需要几分钟，请耐心等待）...")
        
        # 使用包装的嵌入模型，在生成向量时自动更新进度
        progress_embeddings = ProgressEmbeddings(
            embeddings=embeddings,
            progress_callback=progress_callback,
            total_docs=total_docs,
            start_progress=70,
            end_progress=85
        )
        
        # 步骤 4: 创建向量存储（使用带进度更新的嵌入模型）
        if progress_callback:
            progress_callback(85, "🔄 步骤 4/4: 创建向量存储...")
        
        # 兼容不同版本的参数名
        max_retries = 3  # 增加重试次数
        last_error = None
        
        for attempt in range(max_retries):
            try:
                # 在创建前再次确认目录是空的（防止维度不匹配）
                if os.path.exists(db_path):
                    # 检查目录是否为空
                    try:
                        dir_contents = os.listdir(db_path)
                        if dir_contents:
                            # 目录不为空，可能是旧数据库残留，强制清理
                            import shutil
                            if progress_callback:
                                progress_callback(87, f"⚠️ 检测到旧数据库残留，正在清理（尝试 {attempt + 1}/{max_retries}）...")
                            try:
                                shutil.rmtree(db_path)
                                os.makedirs(db_path, exist_ok=True)
                                print(f"✅ 清理了非空数据库目录: {db_path}")
                                import time
                                time.sleep(1)  # 等待文件系统更新
                            except Exception as cleanup_error:
                                print(f"⚠️ 清理目录失败: {str(cleanup_error)}")
                    except Exception:
                        pass  # 如果无法列出目录，继续尝试创建
                
                # 尝试使用 embedding 参数（标准参数名）
                if progress_callback:
                    progress_callback(88, f"🔄 正在创建向量数据库（尝试 {attempt + 1}/{max_retries}）...")
                try:
                    # 优先使用 embedding 参数（这是标准参数名）
                    vectorstore = Chroma.from_documents(
                        documents=documents,
                        embedding=progress_embeddings,
                        persist_directory=db_path
                    )
                    break  # 成功创建，退出循环
                except TypeError as param_error:
                    # 如果 embedding 参数不支持，尝试 embedding_function
                    error_msg = str(param_error).lower()
                    if "embedding" in error_msg or "unexpected keyword" in error_msg:
                        try:
                            vectorstore = Chroma.from_documents(
                                documents=documents,
                                embedding_function=progress_embeddings,
                                persist_directory=db_path
                            )
                            break  # 成功创建，退出循环
                        except TypeError:
                            # 如果两种参数名都不支持，抛出原始错误
                            raise param_error
                    else:
                        raise
            except Exception as create_error:
                # 检查是否是维度不匹配错误
                error_msg = str(create_error).lower()
                if "dimension" in error_msg or "dimensionality" in error_msg:
                    # 维度不匹配，强制清理数据库并重试
                    if progress_callback:
                        progress_callback(87, f"⚠️ 检测到维度不匹配错误，正在清理旧数据库并重试（尝试 {attempt + 1}/{max_retries}）...")
                    import shutil
                    import time
                    try:
                        if os.path.exists(db_path):
                            shutil.rmtree(db_path)
                            print(f"✅ 清理了维度不匹配的数据库目录: {db_path}")
                            time.sleep(2)  # 等待文件系统更新
                        os.makedirs(db_path, exist_ok=True)
                    except Exception as cleanup_error:
                        print(f"⚠️ 清理数据库目录失败: {str(cleanup_error)}")
                    
                    # 如果是最后一次尝试，抛出详细错误
                    if attempt == max_retries - 1:
                        raise Exception(f"创建向量数据库失败：维度不匹配\n\n"
                                      f"错误信息: {str(create_error)}\n\n"
                                      f"可能的原因:\n"
                                      f"- 旧向量数据库使用了不同维度的模型\n"
                                      f"- 模型切换后未正确清理旧数据库\n\n"
                                      f"解决方案:\n"
                                      f"1. 手动删除向量数据库目录: {db_path}\n"
                                      f"2. 或在侧边栏的'向量数据库管理'中删除\n"
                                      f"3. 然后重新加载文件夹") from create_error
                    # 否则继续重试
                    continue
                elif isinstance(create_error, TypeError) and "multiple values for keyword argument" in str(create_error):
                    # 参数重复传递错误，尝试使用不同的参数传递方式
                    if progress_callback:
                        progress_callback(88, f"🔄 检测到参数冲突，尝试使用替代方式创建向量数据库（尝试 {attempt + 1}/{max_retries}）...")
                    try:
                        # 尝试只使用位置参数传递 documents，其他参数使用关键字
                        vectorstore = Chroma.from_documents(
                            documents,
                            embedding=progress_embeddings,
                            persist_directory=db_path
                        )
                        break  # 成功创建，退出循环
                    except Exception as e:
                        # 如果还是失败，记录错误并在最后一次尝试时抛出
                        last_error = e
                        if attempt == max_retries - 1:
                            raise Exception(f"创建向量数据库失败：参数传递错误\n\n"
                                          f"错误信息: {str(create_error)}\n\n"
                                          f"可能的原因:\n"
                                          f"- langchain_chroma 版本不兼容\n"
                                          f"- 参数传递方式冲突\n\n"
                                          f"解决方案:\n"
                                          f"1. 更新 langchain-chroma: pip install --upgrade langchain-chroma\n"
                                          f"2. 检查 Chroma 版本兼容性\n"
                                          f"3. 查看完整错误信息以获取更多细节") from create_error
                        continue
                else:
                    # 其他错误，记录并重试
                    last_error = create_error
                    if attempt == max_retries - 1:
                        raise
                    continue
        else:
            # 所有重试都失败了
            if last_error:
                error_type = type(last_error).__name__
                raise Exception(f"创建向量数据库失败（已重试 {max_retries} 次） [{error_type}]: {str(last_error)}") from last_error
            else:
                raise Exception("创建向量数据库失败：未知错误")
        
        if progress_callback:
            progress_callback(100, "✅ 向量数据库创建完成！")
        
        # 保存文档签名
        if folder_path:
            save_docs_signature(docs_dict, folder_path)
        
        return vectorstore
    except ImportError as e:
        # 导入错误，可能是缺少依赖包
        error_msg = str(e)
        if 'st' in globals():
            st.warning(f"⚠️ 向量数据库功能不可用（缺少依赖包）: {error_msg}\n\n💡 提示：\n- 向量搜索功能将不可用\n- 文档阅读和问答功能仍可正常使用\n- 如需使用向量搜索，请安装: pip install langchain-text-splitters langchain-community chromadb")
        return None
    except Exception as e:
        # 其他错误，显示详细错误信息
        error_msg = str(e)
        error_type = type(e).__name__
        
        # 检查是否是 NumPy 2.0 兼容性问题
        is_numpy_error = ("np.float_" in error_msg or "numpy" in error_msg.lower() or 
                         "AttributeError" in error_type and "float_" in error_msg)
        
        # 记录错误到控制台（用于调试）
        import traceback
        print(f"❌ 向量数据库创建失败:")
        print(f"   错误类型: {error_type}")
        print(f"   错误信息: {error_msg}")
        print(f"   详细堆栈:")
        traceback.print_exc()
        
        # 不再在这里直接显示错误，而是抛出异常让调用者处理
        # 这样调用者可以使用占位符确保错误提示与输入框等宽
        if is_numpy_error:
            # NumPy 2.0 兼容性错误
            error_detail = (f"⚠️ **向量数据库创建失败 - NumPy 版本不兼容**\n\n"
                          f"**错误类型**: `{error_type}`\n\n"
                          f"**错误信息**: {error_msg}\n\n"
                          f"**问题原因**:\n"
                          f"ChromaDB 不兼容 NumPy 2.0，当前环境可能安装了 NumPy 2.0\n\n"
                          f"**解决方案**:\n"
                          f"1. **降级 NumPy 到 1.x 版本**（推荐）:\n"
                          f"   ```bash\n"
                          f"   poetry remove numpy\n"
                          f"   poetry add \"numpy>=1.24.0,<2.0.0\"\n"
                          f"   ```\n"
                          f"   或使用 pip:\n"
                          f"   ```bash\n"
                          f"   pip install \"numpy>=1.24.0,<2.0.0\"\n"
                          f"   ```\n\n"
                          f"2. **重新安装所有依赖**:\n"
                          f"   ```bash\n"
                          f"   poetry install\n"
                          f"   ```\n\n"
                          f"3. **检查 NumPy 版本**:\n"
                          f"   ```bash\n"
                          f"   poetry run python -c \"import numpy; print(numpy.__version__)\"\n"
                          f"   ```\n\n"
                          f"💡 **提示**: 修复后重新运行程序即可")
        else:
            # 其他错误
            error_detail = (f"⚠️ **向量数据库创建失败**\n\n"
                          f"**错误类型**: `{error_type}`\n\n"
                          f"**错误信息**: {error_msg}\n\n"
                          f"**可能的原因**:\n"
                          f"- 文档内容为空或格式不正确\n"
                          f"- 嵌入模型加载失败（检查网络连接或模型文件）\n"
                          f"- ChromaDB 版本不兼容\n"
                          f"- NumPy 版本不兼容（NumPy 2.0 不兼容）\n"
                          f"- 磁盘空间不足或没有写入权限\n"
                          f"- 内存不足（文档太大）\n\n"
                          f"**解决方案**:\n"
                          f"1. 检查文档是否为空\n"
                          f"2. 检查控制台输出的详细错误信息\n"
                          f"3. 尝试减少文档数量或大小\n"
                          f"4. 检查磁盘空间和权限\n"
                          f"5. 如果问题持续，请查看完整错误堆栈\n\n"
                          f"💡 **提示**: 向量搜索功能将不可用，但文档阅读和问答功能仍可正常使用（会使用所有文档内容，可能较慢）")
        
        # 抛出异常，让调用者使用占位符显示错误
        raise Exception(error_detail)

# DeepSeek API接口
def query_deepseek(prompt: str, api_key: str, model: str = "deepseek-chat", max_tokens: int = 2000, 
                   max_retries: int = 3, timeout: int = None):
    """调用DeepSeek API，带重试机制
    
    Args:
        prompt: 提示文本
        api_key: DeepSeek API密钥
        model: 模型名称
        max_tokens: 最大token数
        max_retries: 最大重试次数
        timeout: 超时时间（秒），如果为None则使用默认值或从session_state获取
    """
    import requests
    import time
    
    # 从 session_state 获取超时和重试配置（如果可用）
    if timeout is None:
        if 'st' in globals() and hasattr(st, 'session_state'):
            timeout = st.session_state.get('api_timeout', 60)
        else:
            timeout = 60
    
    if max_retries is None or max_retries == 3:
        if 'st' in globals() and hasattr(st, 'session_state'):
            max_retries = st.session_state.get('api_max_retries', 3)
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": model,
        "messages": [
            {"role": "system", "content": "你是一个有帮助的助手，请基于提供的文档内容回答问题。"},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": 0.3
    }
    
    # 重试机制
    for attempt in range(max_retries):
        try:
            # 根据尝试次数增加超时时间
            current_timeout = timeout + (attempt * 20)
            
            response = requests.post(
                "https://api.deepseek.com/v1/chat/completions",
                headers=headers,
                json=data,
                timeout=current_timeout
            )
            
            if response.status_code == 200:
                result = response.json()
                if "choices" in result and len(result["choices"]) > 0:
                    return result["choices"][0]["message"]["content"]
                else:
                    return "API返回格式异常，请重试"
            elif response.status_code == 401:
                return "API密钥无效，请检查您的DeepSeek API密钥"
            elif response.status_code == 429:
                wait_time = 2 ** attempt  # 指数退避：2秒、4秒、8秒
                if attempt < max_retries - 1:
                    time.sleep(wait_time)
                    continue
                return "API请求频率过高，请稍后再试"
            elif response.status_code == 500:
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)
                    continue
                return "DeepSeek服务器错误，请稍后重试"
            elif response.status_code == 400:
                # 检查是否是上下文长度超限错误
                error_text = response.text.lower()
                if "context" in error_text and ("length" in error_text or "exceeded" in error_text or "too long" in error_text):
                    return "❌ 文档内容过长，超过了API的上下文窗口限制（64K tokens）。\n\n建议：\n1. 减少选择的文档数量\n2. 或者使用分块总结功能（如果可用）\n3. 或者先对每篇文档进行摘要，再总结摘要内容"
                else:
                    return f"API请求参数错误 (状态码: 400): {response.text[:200]}"
            else:
                return f"API请求失败 (状态码: {response.status_code}): {response.text[:200]}"
                
        except requests.exceptions.Timeout:
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                error_msg = f"请求超时（已尝试 {attempt + 1}/{max_retries} 次），{wait_time}秒后重试..."
                if 'st' in globals():
                    st.warning(error_msg)
                time.sleep(wait_time)
                continue
            else:
                return f"请求超时（已重试 {max_retries} 次）。可能的原因：\n1. 网络连接不稳定\n2. 请求内容过长\n3. DeepSeek服务器响应慢\n\n建议：\n- 检查网络连接\n- 尝试减少文档内容\n- 稍后重试"
        
        except requests.exceptions.ConnectionError:
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                error_msg = f"连接错误（已尝试 {attempt + 1}/{max_retries} 次），{wait_time}秒后重试..."
                if 'st' in globals():
                    st.warning(error_msg)
                time.sleep(wait_time)
                continue
            else:
                return "无法连接到DeepSeek API服务器。请检查：\n1. 网络连接是否正常\n2. 是否可以使用代理访问\n3. DeepSeek服务是否正常"
        
        except requests.exceptions.RequestException as e:
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)
                continue
            return f"网络请求异常: {str(e)}"
        
        except Exception as e:
            return f"调用API时出错: {str(e)}"
    
    return "API调用失败，已重试多次仍无法成功"

def search_similar_documents(vectorstore, query: str, k: int = 4):
    """检索相似文档片段"""
    if vectorstore is None:
        return []
    
    try:
        docs = vectorstore.similarity_search(query, k=k)
        return [(doc.page_content, doc.metadata["source"]) for doc in docs]
    except:
        return []

def answer_with_deepseek(question: str, vectorstore, docs_dict: Dict[str, Any], api_key: str):
    """使用DeepSeek回答问题"""
    # 检索相关文档片段
    similar_docs = search_similar_documents(vectorstore, question)
    
    if not similar_docs:
        # 如果没有向量数据库，使用所有文档内容
        context = "\n\n".join([f"文件: {name}\n内容: {data['content'][:2000]}..." 
                             for name, data in docs_dict.items()])
    else:
        # 使用检索到的文档片段
        context_parts = []
        for content, source in similar_docs:
            context_parts.append(f"来自文档 '{source}' 的内容:\n{content}")
        context = "\n\n".join(context_parts)
    
    # 构建提示
    prompt = f"""基于以下文档内容，请回答这个问题：{question}

相关文档内容：
{context[:8000]}  # 限制上下文长度

请基于上述文档内容回答，如果文档中没有相关信息，请明确说明。"""

    return query_deepseek(prompt, api_key)

def generate_summary_deepseek(docs_dict: Dict[str, Any], api_key: str, specific_files: List[str] = None, template_id: str = "default"):
    """使用DeepSeek生成总结报告
    
    Args:
        docs_dict: 文档字典
        api_key: API密钥
        specific_files: 特定文件列表（None表示所有文件）
        template_id: 使用的模版ID（默认为"default"）
    """
    # 提取内容
    contents = []
    if specific_files:
        for filename in specific_files:
            if filename in docs_dict:
                content = docs_dict[filename]['content']
                if isinstance(content, dict):
                    # 移除字段长度限制，让API自行处理
                    content = "\n".join([f"{k}: {v}" for k, v in content.items()])
                contents.append(f"文件: {filename}\n{content}")
    else:
        for filename, data in docs_dict.items():
            content = data['content']
            if isinstance(content, dict):
                # 移除字段长度限制，让API自行处理
                content = "\n".join([f"{k}: {v}" for k, v in content.items()])
            contents.append(f"文件: {filename}\n{content}")
    
    combined_content = "\n\n".join(contents)
    
    # 加载模版
    template_data = get_template("summary", template_id)
    if template_data:
        template_str = template_data.get("template", "")
        # 替换模版中的占位符（移除字符限制，让API自行处理）
        prompt = template_str.format(content=combined_content)
    else:
        # 如果模版不存在，使用默认模版
        prompt = f"""请根据以下文档内容，生成一份详细的总结报告：

文档内容：
{combined_content}

请生成包括以下部分的报告：
1. 整体内容概述
2. 核心要点总结
3. 关键数据/信息提取
4. 主要发现和洞察
5. 建议和下一步行动

报告："""

    return query_deepseek(prompt, api_key, max_tokens=3000)

# 显示版权信息
def show_footer():
    """在页面底部显示版权信息"""
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: #666; padding: 20px 0;'>
            <p style='margin: 5px 0;'><strong>Copyright © 2026 吕滢</strong></p>
            <p style='margin: 5px 0;'>
                GitHub: <a href='https://github.com/lveyond' target='_blank' style='color: #1f77b4;'>@lveyond</a> | 
                QQ/WeChat: 329613507
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

# Streamlit界面
def main():
    # 添加自定义CSS样式，将进度条和primary按钮改为草绿色
    st.markdown("""
    <style>
    /* 进度条颜色改为草绿色 - 使用更通用的选择器 */
    .stProgress .st-bo {
        background-color: #7cb342 !important;
    }
    .stProgress > div > div > div > div {
        background-color: #7cb342 !important;
    }
    div[data-testid="stProgress"] > div > div > div > div {
        background-color: #7cb342 !important;
    }
    
    /* Primary按钮颜色改为草绿色 - 使用更通用的选择器 */
    .stButton > button[kind="primary"],
    .stButton > button[type="primary"],
    button[kind="primary"],
    button[type="primary"] {
        background-color: #7cb342 !important;
        border-color: #7cb342 !important;
        color: white !important;
    }
    
    /* Primary按钮悬停效果 */
    .stButton > button[kind="primary"]:hover,
    .stButton > button[type="primary"]:hover,
    button[kind="primary"]:hover,
    button[type="primary"]:hover {
        background-color: #689f38 !important;
        border-color: #689f38 !important;
    }
    
    /* Primary按钮激活效果 */
    .stButton > button[kind="primary"]:active,
    .stButton > button[type="primary"]:active,
    button[kind="primary"]:active,
    button[type="primary"]:active {
        background-color: #558b2f !important;
        border-color: #558b2f !important;
    }
    
    /* 修复expander组件渲染时短暂显示keyboard_arrow_right文本的问题 */
    /* 这是Streamlit内部Material Icons字体加载时的临时显示问题 */
    /* 确保expander标题区域正确渲染 */
    [data-testid="stExpander"] .streamlit-expanderHeader {
        position: relative;
        overflow: hidden;
    }
    
    [data-testid="stExpander"] .streamlit-expanderHeader label {
        white-space: nowrap;
        overflow: hidden;
        text-overflow: ellipsis;
    }
    
    /* 预加载Material Icons字体，避免显示文本 */
    @import url('https://fonts.googleapis.com/icon?family=Material+Icons');
    
    /* 确保Material Icons正确渲染 */
    [data-testid="stExpander"] [class*="material-icons"],
    [data-testid="stExpander"] .material-icons {
        font-family: 'Material Icons' !important;
        font-feature-settings: 'liga' !important;
        text-rendering: optimizeLegibility !important;
        -webkit-font-smoothing: antialiased !important;
        font-style: normal !important;
        font-weight: normal !important;
        letter-spacing: normal !important;
        text-transform: none !important;
        display: inline-block !important;
        white-space: nowrap !important;
        word-wrap: normal !important;
        direction: ltr !important;
    }
    
    /* 骨架屏样式 */
    .skeleton-screen {
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background: #ffffff;
        z-index: 9999;
        padding: 20px;
        box-sizing: border-box;
    }
    
    .skeleton-header {
        height: 60px;
        background: linear-gradient(90deg, #f0f0f0 25%, #e0e0e0 50%, #f0f0f0 75%);
        background-size: 200% 100%;
        animation: skeleton-loading 1.5s ease-in-out infinite;
        border-radius: 8px;
        margin-bottom: 20px;
        max-width: 600px;
    }
    
    .skeleton-container {
        display: flex;
        gap: 20px;
        height: calc(100vh - 120px);
    }
    
    .skeleton-sidebar {
        width: 300px;
        background: #f8f9fa;
        border-radius: 8px;
        padding: 20px;
    }
    
    .skeleton-sidebar-item {
        height: 40px;
        background: linear-gradient(90deg, #e9ecef 25%, #dee2e6 50%, #e9ecef 75%);
        background-size: 200% 100%;
        animation: skeleton-loading 1.5s ease-in-out infinite;
        border-radius: 6px;
        margin-bottom: 15px;
    }
    
    .skeleton-sidebar-item.short {
        width: 60%;
    }
    
    .skeleton-sidebar-item.medium {
        width: 80%;
    }
    
    .skeleton-main {
        flex: 1;
        background: #ffffff;
        border-radius: 8px;
        padding: 20px;
    }
    
    .skeleton-content-block {
        height: 200px;
        background: linear-gradient(90deg, #f0f0f0 25%, #e0e0e0 50%, #f0f0f0 75%);
        background-size: 200% 100%;
        animation: skeleton-loading 1.5s ease-in-out infinite;
        border-radius: 8px;
        margin-bottom: 20px;
    }
    
    .skeleton-content-line {
        height: 20px;
        background: linear-gradient(90deg, #f0f0f0 25%, #e0e0e0 50%, #f0f0f0 75%);
        background-size: 200% 100%;
        animation: skeleton-loading 1.5s ease-in-out infinite;
        border-radius: 4px;
        margin-bottom: 10px;
    }
    
    .skeleton-content-line.short {
        width: 40%;
    }
    
    .skeleton-content-line.medium {
        width: 70%;
    }
    
    .skeleton-content-line.long {
        width: 100%;
    }
    
    @keyframes skeleton-loading {
        0% {
            background-position: 200% 0;
        }
        100% {
            background-position: -200% 0;
        }
    }
    
    .skeleton-screen.hidden {
        opacity: 0;
        transition: opacity 0.3s ease-out;
        pointer-events: none;
    }
    </style>
    <script>
    // 隐藏expander中可能显示的keyboard_arrow_right文本（Material Icons加载前的临时显示）
    (function() {
        function hideKeyboardArrowText() {
            const expanders = document.querySelectorAll('[data-testid="stExpander"]');
            expanders.forEach(expander => {
                const elements = expander.querySelectorAll('*');
                elements.forEach(element => {
                    const text = element.textContent || element.innerText || '';
                    // 如果元素只包含keyboard_arrow_right文本，隐藏它
                    if (text.trim() === 'keyboard_arrow_right') {
                        element.style.display = 'none';
                        element.style.visibility = 'hidden';
                        element.style.opacity = '0';
                        element.style.width = '0';
                        element.style.height = '0';
                        element.style.overflow = 'hidden';
                    }
                });
            });
        }
        
        // 立即执行
        hideKeyboardArrowText();
        
        // 页面加载完成后执行
        if (document.readyState === 'loading') {
            document.addEventListener('DOMContentLoaded', hideKeyboardArrowText);
        }
        
        // 使用MutationObserver监听DOM变化
        if (typeof MutationObserver !== 'undefined') {
            const observer = new MutationObserver(function(mutations) {
                let shouldCheck = false;
                mutations.forEach(function(mutation) {
                    if (mutation.addedNodes.length > 0) {
                        shouldCheck = true;
                    }
                });
                if (shouldCheck) {
                    setTimeout(hideKeyboardArrowText, 10);
                }
            });
            observer.observe(document.body, {
                childList: true,
                subtree: true
            });
        }
    })();
    
    // 创建并显示骨架屏
    function showSkeletonScreen() {
        const skeleton = document.createElement('div');
        skeleton.className = 'skeleton-screen';
        skeleton.id = 'skeleton-screen';
        skeleton.innerHTML = `
            <div class="skeleton-header"></div>
            <div class="skeleton-container">
                <div class="skeleton-sidebar">
                    <div class="skeleton-sidebar-item short"></div>
                    <div class="skeleton-sidebar-item medium"></div>
                    <div class="skeleton-sidebar-item"></div>
                    <div class="skeleton-sidebar-item short"></div>
                    <div class="skeleton-sidebar-item medium"></div>
                    <div class="skeleton-sidebar-item"></div>
                    <div class="skeleton-sidebar-item short"></div>
                </div>
                <div class="skeleton-main">
                    <div class="skeleton-content-block"></div>
                    <div class="skeleton-content-line long"></div>
                    <div class="skeleton-content-line medium"></div>
                    <div class="skeleton-content-line short"></div>
                    <div class="skeleton-content-block"></div>
                    <div class="skeleton-content-line long"></div>
                    <div class="skeleton-content-line long"></div>
                </div>
            </div>
        `;
        document.body.appendChild(skeleton);
    }
    
    // 隐藏骨架屏
    function hideSkeletonScreen() {
        const skeleton = document.getElementById('skeleton-screen');
        if (skeleton) {
            skeleton.classList.add('hidden');
            setTimeout(() => {
                skeleton.remove();
            }, 300);
        }
    }
    
    // 页面加载时显示骨架屏
    (function() {
        // 立即显示骨架屏
        showSkeletonScreen();
        
        // 检测Streamlit应用是否已加载
        function checkStreamlitLoaded() {
            const stApp = document.querySelector('[data-testid="stApp"]');
            const mainContent = document.querySelector('[data-testid="stAppViewContainer"]');
            
            // 如果Streamlit应用已加载且主要内容已渲染
            if (stApp && mainContent && mainContent.children.length > 0) {
                // 等待一小段时间确保内容完全渲染
                setTimeout(function() {
                    hideSkeletonScreen();
                }, 300);
                return true;
            }
            return false;
        }
        
        // 立即检查一次
        if (!checkStreamlitLoaded()) {
            // 如果还没加载，监听DOM变化
            const observer = new MutationObserver(function(mutations, obs) {
                if (checkStreamlitLoaded()) {
                    obs.disconnect();
                }
            });
            
            observer.observe(document.body, {
                childList: true,
                subtree: true
            });
            
            // 超时保护：最多等待3秒
            setTimeout(function() {
                observer.disconnect();
                hideSkeletonScreen();
            }, 3000);
        }
    })();
    </script>
    """, unsafe_allow_html=True)
    
    st.set_page_config(
        page_title="智能知识库系统 (DeepSeek版)",
        page_icon="📚",
        layout="wide"
    )
    
    st.title("📚 智能知识库系统 (DeepSeek版)")
    st.markdown("---")
    
    # 初始化session state
    if 'docs' not in st.session_state:
        st.session_state.docs = {}
    if 'vectorstore' not in st.session_state:
        st.session_state.vectorstore = None
    if 'selected_file' not in st.session_state:
        st.session_state.selected_file = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'api_key_loaded' not in st.session_state:
        st.session_state.api_key_loaded = False
    if 'is_creating_vectorstore' not in st.session_state:
        st.session_state.is_creating_vectorstore = False
    if 'embedding_model' not in st.session_state:
        st.session_state.embedding_model = load_embedding_model_config()
    
    # 侧边栏
    with st.sidebar:
        st.header("⚙️ 配置")
        
        # DeepSeek API配置
        st.subheader("🔑 DeepSeek API设置")
        
        # 尝试从本地加载 API key
        saved_api_key = None
        if not st.session_state.api_key_loaded:
            saved_api_key = load_api_key()
            if saved_api_key:
                st.session_state.api_key_loaded = True
                st.session_state.saved_api_key = saved_api_key
                st.success("✅ 已从本地加载 API 密钥")
        
        # 使用保存的 key 或用户输入
        default_key = st.session_state.get('saved_api_key', '') if st.session_state.api_key_loaded else ''
        api_key_input = st.text_input(
            "DeepSeek API密钥", 
            value=default_key,
            type="password", 
            help="在 https://platform.deepseek.com/ 获取API密钥\n💾 密钥会自动保存到本地，下次启动无需重新输入"
        )
        
        # API key 管理按钮
        col_key1, col_key2 = st.columns(2)
        with col_key1:
            if st.button("💾 保存密钥", use_container_width=True):
                if api_key_input:
                    if save_api_key(api_key_input):
                        st.session_state.saved_api_key = api_key_input
                        st.session_state.api_key_loaded = True
                        st.success("✅ API 密钥已保存到本地")
                        st.rerun()
                else:
                    st.warning("请输入 API 密钥")
        
        with col_key2:
            if st.button("🗑️ 清除密钥", use_container_width=True):
                if delete_api_key():
                    st.session_state.saved_api_key = ""
                    st.session_state.api_key_loaded = False
                    st.success("✅ 已清除本地保存的 API 密钥")
                    st.rerun()
        
        # 自动保存逻辑：如果用户输入了新密钥且与保存的不同，自动保存
        if api_key_input and api_key_input != st.session_state.get('saved_api_key', ''):
            # 用户输入了新密钥，自动保存（仅在首次输入时，静默保存）
            if not st.session_state.api_key_loaded or api_key_input != saved_api_key:
                if save_api_key(api_key_input, show_error=False):
                    st.session_state.saved_api_key = api_key_input
                    st.session_state.api_key_loaded = True
                    # 静默保存，不显示提示（避免频繁刷新）
        
        # 使用输入的 key（优先使用新输入的）
        api_key = api_key_input if api_key_input else (st.session_state.get('saved_api_key', '') if st.session_state.api_key_loaded else '')
        
        # 显示保存状态
        if os.path.exists(CONFIG_FILE):
            saved_time = ""
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    saved_time = config.get("saved_at", "")
            except:
                pass
            if saved_time:
                st.caption(f"💾 密钥已保存（{saved_time}）")
            else:
                st.caption("💾 密钥已保存到本地")
        else:
            st.caption("💡 输入密钥后会自动保存到本地")
        
        # 模型选择
        model_choice = st.selectbox(
            "选择模型",
            ["deepseek-chat", "deepseek-coder"],
            help="deepseek-chat: 通用对话模型\ndeepseek-coder: 代码专用模型"
        )
        
        # API 超时和重试配置（高级设置）
        if 'api_timeout' not in st.session_state:
            st.session_state.api_timeout = 60
        if 'api_max_retries' not in st.session_state:
            st.session_state.api_max_retries = 3
        
        # API 超时和重试配置（高级设置）
        with st.expander("⚙️ 高级设置（网络问题时可调整）", expanded=False):
            timeout_seconds = st.slider(
                "请求超时时间（秒）",
                min_value=30,
                max_value=180,
                value=st.session_state.api_timeout,
                step=10,
                help="如果经常超时，可以增加此值"
            )
            max_retries = st.slider(
                "最大重试次数",
                min_value=1,
                max_value=5,
                value=st.session_state.api_max_retries,
                step=1,
                help="网络不稳定时可以增加重试次数"
            )
            
            # 保存到 session state
            st.session_state.api_timeout = timeout_seconds
            st.session_state.api_max_retries = max_retries
        
        st.markdown("---")
        
        # 嵌入模型配置
        st.subheader("🤖 嵌入模型设置")
        
        # 初始化嵌入模型配置
        if 'embedding_model' not in st.session_state:
            st.session_state.embedding_model = load_embedding_model_config()
        
        # 可用的嵌入模型列表
        embedding_models = {
            "BAAI/bge-small-zh-v1.5": {
                "name": "bge-small-zh-v1.5",
                "description": "轻量快速（384维，~130MB）",
                "size": "~130MB",
                "performance": "⭐⭐⭐"
            },
            "BAAI/bge-base-zh-v1.5": {
                "name": "bge-base-zh-v1.5",
                "description": "平衡性能（768维，~420MB）",
                "size": "~420MB",
                "performance": "⭐⭐⭐⭐"
            },
            "BAAI/bge-large-zh-v1.5": {
                "name": "bge-large-zh-v1.5",
                "description": "最佳性能（1024维，~1.2GB）",
                "size": "~1.2GB",
                "performance": "⭐⭐⭐⭐⭐"
            }
        }
        
        # 模型选择下拉框
        model_options = [f"{info['name']} - {info['description']}" for model_id, info in embedding_models.items()]
        current_model_index = 0
        for idx, (model_id, info) in enumerate(embedding_models.items()):
            if model_id == st.session_state.embedding_model:
                current_model_index = idx
                break
        
        selected_model_display = st.selectbox(
            "选择嵌入模型",
            options=model_options,
            index=current_model_index,
            help="用于文档向量化的模型。更大的模型性能更好，但需要更多内存和存储空间。"
        )
        
        # 获取选中的模型ID
        selected_model_id = list(embedding_models.keys())[model_options.index(selected_model_display)]
        
        # 显示当前模型信息
        current_model_info = embedding_models[selected_model_id]
        st.caption(f"📊 性能: {current_model_info['performance']} | 💾 大小: {current_model_info['size']}")
        
        # 检查模型是否已下载
        model_path = get_model_path(selected_model_id)
        # 检查模型是否存在（本地路径或HuggingFace缓存）
        cache_dir = os.path.join(os.path.expanduser("~"), ".cache", "huggingface", "hub")
        cache_model_name = f"models--{selected_model_id.replace('/', '--')}"
        cache_path = os.path.join(cache_dir, cache_model_name)
        
        # 判断模型是否存在
        model_exists = (
            (os.path.exists(model_path) and os.path.isdir(model_path)) or  # 本地路径存在
            os.path.exists(cache_path)  # HuggingFace缓存存在
        )
        
        # 如果模型不存在，显示下载选项
        if not model_exists and selected_model_id != "BAAI/bge-small-zh-v1.5":
            if not os.path.exists(cache_path):
                with st.expander("📥 下载模型", expanded=False):
                    st.info(f"模型 {selected_model_id} 尚未下载，首次使用需要下载。")
                    if st.button(f"⬇️ 下载 {current_model_info['name']}", use_container_width=True, key=f"download_{selected_model_id}"):
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        success = download_model(
                            selected_model_id,
                            progress_callback=lambda p, msg: (
                                progress_bar.progress(p / 100.0),
                                status_text.text(msg)
                            )
                        )
                        
                        if success:
                            st.success(f"✅ 模型 {current_model_info['name']} 下载完成！")
                            st.info("💡 请刷新页面后使用新模型")
                        else:
                            st.error("❌ 下载失败，请检查网络连接")
                        
                        import time
                        time.sleep(1)
                        progress_bar.empty()
                        status_text.empty()
            else:
                st.success(f"✅ 模型 {current_model_info['name']} 已下载")
        
        # 保存模型选择
        if selected_model_id != st.session_state.embedding_model:
            # 检查是否有现有的向量数据库
            has_existing_db = False
            if os.path.exists("./chroma_db"):
                try:
                    db_dirs = [d for d in os.listdir("./chroma_db") 
                              if os.path.isdir(os.path.join("./chroma_db", d))]
                    has_existing_db = len(db_dirs) > 0
                except:
                    pass
            
            if save_embedding_model_config(selected_model_id):
                st.session_state.embedding_model = selected_model_id
                st.success(f"✅ 已切换到 {current_model_info['name']}")
                
                if has_existing_db:
                    st.warning("⚠️ **重要提示**：切换模型后，现有的向量数据库将无法使用（维度不匹配）")
                    st.info("💡 **操作建议**：\n"
                           "1. 切换模型后，系统会在下次创建向量数据库时自动清理旧数据库\n"
                           "2. 或者手动删除向量数据库：在侧边栏的'向量数据库管理'中删除\n"
                           "3. 然后重新加载文件夹或上传文件，系统会使用新模型重新创建向量数据库")
                else:
                    st.info("💡 切换模型后，下次创建向量数据库时将使用新模型")
                st.rerun()
        
        st.markdown("---")
        st.header("📁 文件管理")
        
        # 文件夹选择
        folder_path = st.text_input("文件夹路径", placeholder="输入文件夹路径，如: ./documents")
        
        # 在列布局外创建占位符，确保与输入框等宽
        info_placeholder = st.empty()
        progress_placeholder = st.empty()
        status_placeholder = st.empty()
        success_placeholder = st.empty()  # 用于显示成功消息，确保等宽
        error_placeholder = st.empty()  # 用于显示错误消息，确保等宽
        
        col1, col2 = st.columns(2)
        with col1:
            # 检查是否正在创建向量数据库
            is_creating_vectorstore = st.session_state.get('is_creating_vectorstore', False)
            
            if st.button("📂 加载文件夹", use_container_width=True, disabled=is_creating_vectorstore):
                if folder_path and os.path.exists(folder_path) and os.path.isdir(folder_path):
                    with st.spinner("正在读取文件..."):
                        st.session_state.docs = process_folder(folder_path)
                        # 保存当前文件夹路径
                        st.session_state.current_folder_path = folder_path
                    
                    # 显示已加载文件信息（在列布局外，确保与输入框等宽）
                    if st.session_state.docs:
                        info_placeholder.info(f"📄 已加载 {len(st.session_state.docs)} 个文件")
                    
                    # 检查并加载/创建向量数据库
                    if st.session_state.docs:
                        st.session_state.is_creating_vectorstore = True
                        # 进度条和状态文本（在列布局外，确保与输入框等宽）
                        progress_bar = progress_placeholder.progress(0)
                        status_text = status_placeholder.empty()
                        
                        try:
                            # 首先尝试加载已有向量数据库
                            status_text.text("🔄 正在检查已有向量数据库...")
                            progress_bar.progress(0.05)
                            
                            existing_vectorstore = load_existing_vector_store(
                                folder_path=folder_path,
                                progress_callback=lambda p, msg: (
                                    progress_bar.progress(p / 100.0),
                                    status_text.text(msg)
                                )
                            )
                            
                            # 检查文档是否变化
                            docs_changed = check_docs_changed(st.session_state.docs, folder_path)
                            
                            # 根据文档变化和数据库加载情况决定操作
                            if not docs_changed:
                                # 文档未变化
                                if existing_vectorstore:
                                    # 数据库可用，使用已有向量数据库
                                    st.session_state.vectorstore = existing_vectorstore
                                    progress_bar.progress(1.0)
                                    status_text.text("✅ 已加载已有向量数据库！")
                                    success_placeholder.success("✅ 已加载已有向量数据库（文档未变化）")
                                else:
                                    # 文档未变化但数据库无法加载（可能损坏）
                                    # 不自动重新创建，提示用户
                                    progress_bar.progress(1.0)
                                    status_text.text("⚠️ 文档未变化，但向量数据库无法加载")
                                    warning_msg = (
                                        "⚠️ **文档未变化，但向量数据库无法加载**\n\n"
                                        "可能的原因：\n"
                                        "- 向量数据库文件损坏\n"
                                        "- 数据库版本不兼容\n\n"
                                        "**建议**：\n"
                                        "- 如果数据库损坏，可以手动删除数据库目录后重新创建\n"
                                        "- 或者点击'重新加载'按钮强制重新创建"
                                    )
                                    error_placeholder.warning(warning_msg)
                                    st.session_state.vectorstore = None
                            else:
                                # 文档已变化，需要重新创建
                                if existing_vectorstore:
                                    status_text.text("📝 检测到文档变化，正在重新创建向量数据库...")
                                    progress_bar.progress(0.01)
                                
                                # 直接调用 create_local_vector_store，让它自己管理所有进度更新
                                try:
                                    st.session_state.vectorstore = create_local_vector_store(
                                        st.session_state.docs,
                                        folder_path=folder_path,
                                        progress_callback=lambda p, msg: (
                                            progress_bar.progress(p / 100.0),
                                            status_text.text(msg)
                                        )
                                    )
                                    
                                    if st.session_state.vectorstore:
                                        progress_bar.progress(1.0)
                                        status_text.text("✅ 向量数据库创建完成！")
                                        success_placeholder.success("✅ 向量数据库创建完成！")
                                        error_placeholder.empty()  # 清空错误提示
                                    else:
                                        progress_placeholder.empty()
                                        status_placeholder.empty()
                                        # 如果创建失败，已经在 create_local_vector_store 中显示了警告
                                except Exception as create_error:
                                    # 捕获异常，使用占位符显示错误（确保等宽）
                                    progress_placeholder.empty()
                                    status_placeholder.empty()
                                    error_type = type(create_error).__name__
                                    error_msg = str(create_error)
                                    error_placeholder.error(f"⚠️ **向量数据库创建失败**\n\n"
                                                          f"**错误类型**: `{error_type}`\n\n"
                                                          f"**错误信息**: {error_msg}")
                                    st.session_state.vectorstore = None
                        finally:
                            st.session_state.is_creating_vectorstore = False
                            # 清理进度条
                            import time
                            time.sleep(0.5)
                            progress_placeholder.empty()
                            status_placeholder.empty()
                else:
                    st.error("请输入有效的文件夹路径")
        
        with col2:
            if st.button("🔄 重新加载", use_container_width=True, disabled=is_creating_vectorstore):
                # 获取当前文件夹路径
                current_folder_path = st.session_state.get('current_folder_path', None)
                
                if current_folder_path and os.path.exists(current_folder_path) and os.path.isdir(current_folder_path):
                    # 如果有当前文件夹路径，重新加载并强制重新创建向量数据库
                    st.session_state.is_creating_vectorstore = True
                    progress_bar = progress_placeholder.progress(0)
                    status_text = status_placeholder.empty()
                    
                    try:
                        # 重新读取文件
                        with st.spinner("正在重新读取文件..."):
                            st.session_state.docs = process_folder(current_folder_path)
                        
                        if st.session_state.docs:
                            info_placeholder.info(f"📄 已加载 {len(st.session_state.docs)} 个文件")
                            
                            # 强制重新创建向量数据库（即使文档未变化）
                            status_text.text("🔄 正在强制重新创建向量数据库...")
                            progress_bar.progress(0.01)
                            
                            # 删除旧的数据库目录（如果存在）
                            db_path = get_vector_db_path(current_folder_path)
                            if os.path.exists(db_path):
                                try:
                                    import shutil
                                    shutil.rmtree(db_path)
                                    print(f"[INFO] 已删除旧数据库目录: {db_path}")
                                except Exception as e:
                                    print(f"[WARN] 删除旧数据库目录失败: {str(e)}")
                            
                            # 创建新的向量数据库
                            try:
                                st.session_state.vectorstore = create_local_vector_store(
                                    st.session_state.docs,
                                    folder_path=current_folder_path,
                                    progress_callback=lambda p, msg: (
                                        progress_bar.progress(p / 100.0),
                                        status_text.text(msg)
                                    )
                                )
                                
                                if st.session_state.vectorstore:
                                    progress_bar.progress(1.0)
                                    status_text.text("✅ 向量数据库重新创建完成！")
                                    success_placeholder.success("✅ 向量数据库重新创建完成！")
                                    error_placeholder.empty()
                                else:
                                    progress_placeholder.empty()
                                    status_placeholder.empty()
                                    error_placeholder.error("❌ 向量数据库创建失败")
                            except Exception as create_error:
                                progress_placeholder.empty()
                                status_placeholder.empty()
                                error_type = type(create_error).__name__
                                error_msg = str(create_error)
                                error_placeholder.error(f"⚠️ **向量数据库创建失败**\n\n"
                                                      f"**错误类型**: `{error_type}`\n\n"
                                                      f"**错误信息**: {error_msg}")
                                st.session_state.vectorstore = None
                    finally:
                        st.session_state.is_creating_vectorstore = False
                        import time
                        time.sleep(0.5)
                        progress_placeholder.empty()
                        status_placeholder.empty()
                else:
                    # 如果没有当前文件夹路径，只清空状态
                    st.session_state.docs = {}
                    st.session_state.vectorstore = None
                    st.session_state.is_creating_vectorstore = False
                    info_placeholder.empty()
                    progress_placeholder.empty()
                    status_placeholder.empty()
                    success_placeholder.empty()
                    error_placeholder.empty()
                    st.rerun()
        
        # 如果不在创建过程中，显示已加载文件信息
        if st.session_state.get('docs') and not is_creating_vectorstore:
            info_placeholder.info(f"📄 已加载 {len(st.session_state.docs)} 个文件")
        
        # 文件上传
        st.subheader("或上传文件")
        uploaded_files = st.file_uploader(
            "选择文件",
            type=['txt', 'docx', 'pdf', 'xlsx', 'xls', 'md', 'js', 'json'],
            accept_multiple_files=True,
            label_visibility="collapsed"
        )
        
        # 在文件上传区域创建占位符，确保与上传组件等宽
        upload_info_placeholder = st.empty()
        upload_progress_placeholder = st.empty()
        upload_status_placeholder = st.empty()
        upload_success_placeholder = st.empty()  # 用于显示成功消息，确保等宽
        upload_error_placeholder = st.empty()  # 用于显示错误消息，确保等宽
        
        if uploaded_files and st.button("上传文件"):
            temp_dir = tempfile.mkdtemp()
            for uploaded_file in uploaded_files:
                file_path = os.path.join(temp_dir, uploaded_file.name)
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                
                # 读取文件
                filename = uploaded_file.name
                file_ext = filename.split('.')[-1].lower()
                
                try:
                    if file_ext == 'txt':
                        content = read_text_file(file_path)
                    elif file_ext == 'docx':
                        content = read_docx_file(file_path)
                    elif file_ext == 'pdf':
                        content = read_pdf_file(file_path)
                    elif file_ext in ['xlsx', 'xls']:
                        content = read_excel_file(file_path)
                    elif file_ext == 'md':
                        content = read_markdown_file(file_path)
                    elif file_ext == 'js':
                        content = read_javascript_file(file_path)
                    elif file_ext == 'json':
                        content = read_json_file(file_path)
                    else:
                        content = f"不支持的文件类型: {file_ext}"
                    
                    st.session_state.docs[filename] = {
                        'path': file_path,
                        'content': content,
                        'type': file_ext,
                        'size': uploaded_file.size
                    }
                except Exception as e:
                    st.error(f"读取文件 {filename} 失败: {str(e)}")
            
            # 显示已上传文件信息（在占位符中，确保与上传组件等宽）
            if uploaded_files:
                upload_info_placeholder.info(f"📄 已上传 {len(uploaded_files)} 个文件")
            
            # 检查并加载/创建向量数据库（上传文件时 folder_path 为 None）
            if st.session_state.docs:
                st.session_state.is_creating_vectorstore = True
                # 进度条和状态文本（在占位符中，确保与上传组件等宽）
                progress_bar = upload_progress_placeholder.progress(0)
                status_text = upload_status_placeholder.empty()
                
                try:
                    # 首先尝试加载已有向量数据库
                    status_text.text("🔄 正在检查已有向量数据库...")
                    progress_bar.progress(0.05)
                    
                    existing_vectorstore = load_existing_vector_store(
                        folder_path=None,  # 上传文件时没有文件夹路径
                        progress_callback=lambda p, msg: (
                            progress_bar.progress(p / 100.0),
                            status_text.text(msg)
                        )
                    )
                    
                    # 检查文档是否变化
                    docs_changed = check_docs_changed(st.session_state.docs, None)
                    
                    # 如果 load_existing_vector_store 返回了 vectorstore，说明数据库可用
                    # 不需要再次验证，避免重复验证导致的问题
                    if existing_vectorstore and not docs_changed:
                        # 文档未变化，使用已有向量数据库
                        st.session_state.vectorstore = existing_vectorstore
                        progress_bar.progress(1.0)
                        status_text.text("✅ 已加载已有向量数据库！")
                        upload_success_placeholder.success("✅ 已加载已有向量数据库（文档未变化）")
                    else:
                        # 文档变化或不存在，需要重新创建
                        # 移除所有硬编码的进度更新，让 create_local_vector_store 完全控制进度
                        if existing_vectorstore and docs_changed:
                            status_text.text("📝 检测到文档变化，正在重新创建向量数据库...")
                            progress_bar.progress(0.01)
                        
                        # 直接调用 create_local_vector_store，让它自己管理所有进度更新
                        try:
                            st.session_state.vectorstore = create_local_vector_store(
                                st.session_state.docs,
                                folder_path=None,  # 上传文件时没有文件夹路径
                                progress_callback=lambda p, msg: (
                                    progress_bar.progress(p / 100.0),
                                    status_text.text(msg)
                                )
                            )
                            
                            if st.session_state.vectorstore:
                                progress_bar.progress(1.0)
                                status_text.text("✅ 向量数据库更新完成！")
                                upload_success_placeholder.success("✅ 向量数据库更新完成！")
                                upload_error_placeholder.empty()  # 清空错误提示
                            else:
                                upload_progress_placeholder.empty()
                                upload_status_placeholder.empty()
                        except Exception as create_error:
                            # 捕获异常，使用占位符显示错误（确保等宽）
                            upload_progress_placeholder.empty()
                            upload_status_placeholder.empty()
                            error_msg = str(create_error)
                            upload_error_placeholder.error(error_msg)
                            st.session_state.vectorstore = None
                finally:
                    st.session_state.is_creating_vectorstore = False
                    import time
                    time.sleep(0.5)
                    upload_progress_placeholder.empty()
                    upload_status_placeholder.empty()
        
        # 如果不在创建过程中，显示已上传文件信息
        if st.session_state.get('docs') and not st.session_state.get('is_creating_vectorstore', False) and uploaded_files:
            upload_info_placeholder.info(f"📄 已上传 {len(uploaded_files)} 个文件")
        
        # 如果不在创建过程中，显示已上传文件信息
        if st.session_state.get('docs') and not st.session_state.get('is_creating_vectorstore', False) and uploaded_files:
            upload_info_placeholder.info(f"📄 已上传 {len(uploaded_files)} 个文件")
        
        # 文件统计
        if st.session_state.docs:
            st.markdown("---")
            st.header("📊 统计信息")
            total_files = len(st.session_state.docs)
            file_types = {}
            for data in st.session_state.docs.values():
                file_type = data['type']
                file_types[file_type] = file_types.get(file_type, 0) + 1
            
            st.write(f"**文件总数**: {total_files}")
            for ftype, count in file_types.items():
                st.write(f"**{ftype}文件**: {count}个")
        
        st.markdown("---")
        
        # 向量数据库管理
        if st.session_state.docs:
            with st.expander("🗑️ 向量数据库管理", expanded=False):
                col_vdb1, col_vdb2 = st.columns(2)
                
                with col_vdb1:
                    # 获取当前文件夹的向量数据库路径
                    current_folder_path = st.session_state.get('current_folder_path', None)
                    if current_folder_path:
                        current_db_path = get_vector_db_path(current_folder_path)
                        if st.button("🗑️ 删除当前向量数据库", use_container_width=True):
                            import shutil
                            if os.path.exists(current_db_path):
                                try:
                                    shutil.rmtree(current_db_path)
                                    st.session_state.vectorstore = None
                                    st.success("✅ 当前向量数据库已删除")
                                    st.info("💡 下次加载相同文件夹时会自动重新创建")
                                except Exception as e:
                                    st.error(f"删除失败: {str(e)}")
                            else:
                                st.info("当前向量数据库不存在")
                    else:
                        if st.button("🗑️ 删除所有向量数据库", use_container_width=True):
                            import shutil
                            if os.path.exists("./chroma_db"):
                                try:
                                    shutil.rmtree("./chroma_db")
                                    st.session_state.vectorstore = None
                                    st.success("✅ 所有向量数据库已删除")
                                    st.info("💡 下次加载文档时会自动重新创建")
                                except Exception as e:
                                    st.error(f"删除失败: {str(e)}")
                            else:
                                st.info("向量数据库不存在")
                
                with col_vdb2:
                    # 显示所有向量数据库信息
                    if os.path.exists("./chroma_db"):
                        try:
                            import shutil
                            total_size = 0
                            db_count = 0
                            for item in os.listdir("./chroma_db"):
                                item_path = os.path.join("./chroma_db", item)
                                if os.path.isdir(item_path):
                                    db_count += 1
                                    total_size += sum(f.stat().st_size for f in Path(item_path).rglob('*') if f.is_file())
                            
                            size_mb = total_size / (1024 * 1024)
                            if db_count > 0:
                                st.caption(f"💾 共 {db_count} 个向量数据库")
                                st.caption(f"💾 总大小: {size_mb:.2f} MB")
                            else:
                                st.caption("💾 向量数据库存在")
                        except:
                            st.caption("💾 向量数据库存在")
                    else:
                        st.caption("💾 未创建向量数据库")
                
                # 显示 HuggingFace 缓存信息
                hf_cache_path = os.path.join(os.path.expanduser("~"), ".cache", "huggingface")
                if os.path.exists(hf_cache_path):
                    try:
                        hf_size = sum(f.stat().st_size for f in Path(hf_cache_path).rglob('*') if f.is_file())
                        hf_size_gb = hf_size / (1024 * 1024 * 1024)
                        st.caption(f"🤖 HuggingFace 模型缓存: {hf_size_gb:.2f} GB")
                        st.caption(f"   位置: {hf_cache_path}")
                        st.caption("   💡 如需清理，可手动删除该目录")
                    except:
                        pass
        
        st.caption("💡 提示：使用本地向量数据库进行语义搜索，无需API密钥")
    
    # 主界面
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📄 文档浏览器")
        
        if st.session_state.docs:
            # 文件列表
            files = list(st.session_state.docs.keys())
            selected_file = st.selectbox(
                "选择文件浏览",
                files,
                index=0,
                key="file_selector"
            )
            
            if selected_file:
                file_info = st.session_state.docs[selected_file]
                
                # 文件信息卡片
                st.info(f"""
                **文件信息**
                - 类型: {file_info['type'].upper()}
                - 大小: {file_info['size']:,} 字节
                """)
                
                # 内容显示
                with st.expander("📋 查看内容", expanded=True):
                    content = file_info['content']
                    if isinstance(content, dict):  # Excel文件
                        tab_names = list(content.keys())
                        tabs = st.tabs(tab_names)
                        for i, (sheet_name, sheet_content) in enumerate(content.items()):
                            with tabs[i]:
                                st.text_area(
                                    f"{sheet_name} 工作表",
                                    sheet_content,
                                    height=300,
                                    key=f"excel_{selected_file}_{i}"
                                )
                    else:
                        st.text_area(
                            "文件内容",
                            content,
                            height=400,
                            key=f"content_{selected_file}"
                        )
        
        # 批量操作
        if st.session_state.docs:
            st.markdown("---")
            st.subheader("📈 批量总结")
            
            # 文档选择选项
            summary_mode = st.radio(
                "选择总结范围",
                ["📚 所有文档", "📄 选择特定文档"],
                horizontal=True,
                help="选择要总结的文档范围"
            )
            
            selected_files_for_summary = []
            if summary_mode == "📄 选择特定文档":
                # 多选文档
                file_options = list(st.session_state.docs.keys())
                selected_files_for_summary = st.multiselect(
                    "选择要总结的文档（可多选）",
                    options=file_options,
                    default=[],
                    help="选择要包含在总结报告中的文档"
                )
                
                if not selected_files_for_summary:
                    st.info("💡 请至少选择一个文档")
                    summary_button_disabled = True
                else:
                    summary_button_disabled = False
                    st.info(f"✅ 已选择 {len(selected_files_for_summary)} 个文档")
            else:
                summary_button_disabled = False
                st.info(f"📚 将总结所有 {len(st.session_state.docs)} 个文档")
            
            # Prompt模版选择和管理
            # 加载模版列表
            summary_templates = load_templates("summary")
            template_options = {f"{t['name']} ({tid})": tid for tid, t in summary_templates.items()}
            
            # 初始化session state
            if 'selected_summary_template' not in st.session_state:
                st.session_state.selected_summary_template = "default"
            
            col_template1, col_template1_btn, col_template2, col_template3 = st.columns([3, 1, 1, 1])
            with col_template1:
                selected_template_display = st.selectbox(
                    "选择Prompt模版",
                    options=list(template_options.keys()),
                    index=list(template_options.values()).index(st.session_state.selected_summary_template) if st.session_state.selected_summary_template in template_options.values() else 0,
                    help="选择用于生成总结的Prompt模版"
                )
                st.session_state.selected_summary_template = template_options[selected_template_display]
            
            with col_template1_btn:
                # 刷新按钮
                st.markdown("<div style='height: 27px;'></div>", unsafe_allow_html=True)
                if st.button("🔄 刷新", use_container_width=True, key="refresh_summary_template_btn", help="刷新模板列表以获取最新模板"):
                    # 重新加载模板
                    summary_templates = load_templates("summary")
                    template_options = {f"{t['name']} ({tid})": tid for tid, t in summary_templates.items()}
                    # 如果当前选中的模板不存在，重置为default
                    if st.session_state.selected_summary_template not in template_options.values():
                        st.session_state.selected_summary_template = "default"
                    st.rerun()
            
            with col_template2:
                # 添加占位符以对齐selectbox的label和help icon
                st.markdown("<div style='height: 27px;'></div>", unsafe_allow_html=True)
                if st.button("👁️ 预览", use_container_width=True, key="preview_summary_template_btn"):
                    st.session_state.show_summary_template_preview = True
            
            with col_template3:
                # 添加占位符以对齐selectbox的label和help icon
                st.markdown("<div style='height: 27px;'></div>", unsafe_allow_html=True)
                # 检查选中的模版是否可以删除（默认模版不可删除）
                can_delete = not is_default_template("summary", st.session_state.selected_summary_template)
                if st.button("🗑️ 删除", use_container_width=True, key="delete_summary_template_btn", disabled=not can_delete):
                    if delete_template("summary", st.session_state.selected_summary_template):
                        st.success("✅ 模版已删除")
                        st.session_state.selected_summary_template = "default"
                        st.rerun()
            
            # 模版预览
            if st.session_state.get('show_summary_template_preview', False):
                template_data = get_template("summary", st.session_state.selected_summary_template)
                if template_data:
                    with st.expander("📋 模版预览", expanded=True):
                        st.markdown(f"**模版名称**: {template_data.get('name', '')}")
                        st.markdown(f"**模版描述**: {template_data.get('description', '')}")
                        st.markdown("**模版内容**:")
                        st.code(template_data.get('template', ''), language='text')
                        if st.button("关闭预览", key="close_preview_summary"):
                            st.session_state.show_summary_template_preview = False
            
            # 模版管理（仅保留创建/编辑功能）
            with st.expander("⚙️ 模版管理", expanded=False):
                st.markdown("**创建/编辑模版**")
                new_template_name = st.text_input("模版名称", key="new_summary_template_name")
                new_template_desc = st.text_input("模版描述", key="new_summary_template_desc")
                new_template_content = st.text_area(
                    "模版内容（使用 {content} 作为文档内容占位符）",
                    height=200,
                    key="new_summary_template_content",
                    help="示例：请根据以下文档内容，生成总结：\n\n文档内容：\n{content}\n\n总结："
                )
                if st.button("💾 保存模版", key="save_summary_template"):
                    if new_template_name and new_template_content:
                        if save_template("summary", "", new_template_name, new_template_desc, new_template_content):
                            st.success("✅ 模版已保存")
                            st.rerun()
                    else:
                        st.warning("请输入模版名称和内容")
            
            # 生成报告按钮
            generate_summary_clicked = st.button(
                "生成知识库总结报告", 
                use_container_width=True,
                disabled=summary_button_disabled if summary_mode == "📄 选择特定文档" else False
            )
            
            # 处理生成总结报告的逻辑
            if generate_summary_clicked:
                if not api_key:
                    st.error("请先输入DeepSeek API密钥")
                else:
                    # 确定要总结的文档
                    if summary_mode == "📚 所有文档":
                        files_to_summarize = None  # None 表示所有文档
                        summary_title = f"所有文档总结（共 {len(st.session_state.docs)} 个文档）"
                    else:
                        files_to_summarize = selected_files_for_summary
                        summary_title = f"选定文档总结（共 {len(selected_files_for_summary)} 个文档）"
                    
                    # 获取选中的模版名称用于显示
                    selected_template_data = get_template("summary", st.session_state.selected_summary_template)
                    template_name = selected_template_data.get('name', '默认模版') if selected_template_data else '默认模版'
                    summary_title += f" - {template_name}"
                    
                    with st.spinner(f"正在生成总结报告（{summary_title}）..."):
                        summary = generate_summary_deepseek(
                            st.session_state.docs, 
                            api_key,
                            specific_files=files_to_summarize,
                            template_id=st.session_state.selected_summary_template
                        )
                        # 生成时间戳（在生成总结时生成，确保同一总结使用相同时间戳）
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        # 保存到 session state
                        st.session_state.summary = summary
                        st.session_state.summary_title = summary_title
                        st.session_state.summary_timestamp = timestamp  # 保存时间戳
                        # 保存文档列表信息用于后续显示
                        if files_to_summarize:
                            st.session_state.summary_files = files_to_summarize
                            st.session_state.summary_doc_count = len(files_to_summarize)
                        else:
                            st.session_state.summary_files = None
                            st.session_state.summary_doc_count = len(st.session_state.docs)
            
            # 显示总结报告（移到按钮if块外，确保始终显示在col1）
            if 'summary' in st.session_state and st.session_state.summary:
                st.markdown("---")
                summary_title_display = st.session_state.get('summary_title', '总结报告')
                with st.expander(f"📊 查看总结报告 - {summary_title_display}", expanded=True):
                    # 显示总结的文档信息
                    summary_files = st.session_state.get('summary_files', None)
                    summary_doc_count = st.session_state.get('summary_doc_count', 0)
                    if summary_files:
                        st.markdown(f"**总结的文档：** {', '.join(summary_files)}")
                    else:
                        st.markdown(f"**总结的文档：** 所有文档（共 {summary_doc_count} 个）")
                    st.markdown("---")
                    st.write(st.session_state.summary)
                
                # 保存报告选项
                st.markdown("#### 💾 保存报告")
                col_save1, col_save2, col_save3 = st.columns(3)
                
                # 使用生成总结时保存的时间戳，如果没有则生成新的（兼容旧代码）
                timestamp = st.session_state.get('summary_timestamp', datetime.now().strftime("%Y%m%d_%H%M%S"))
                
                # 确定文档列表用于保存
                summary_files = st.session_state.get('summary_files', None)
                summary_doc_count = st.session_state.get('summary_doc_count', 0)
                if summary_files:
                    doc_list = summary_files
                    doc_count = len(summary_files)
                else:
                    doc_list = list(st.session_state.docs.keys())
                    doc_count = summary_doc_count
                
                # 生成 Markdown 格式内容
                md_summary = f"# {summary_title_display}\n\n生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n文档数量: {doc_count}\n\n---\n\n{st.session_state.summary}"
                
                with col_save1:
                    st.download_button(
                        label="📄 保存为TXT",
                        data=st.session_state.summary,
                        file_name=f"知识库总结_{timestamp}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                
                with col_save2:
                    st.download_button(
                        label="📝 保存为Markdown",
                        data=md_summary,
                        file_name=f"知识库总结_{timestamp}.md",
                        mime="text/markdown",
                        use_container_width=True
                    )
                
                with col_save3:
                    # JSON格式（包含元数据）
                    json_data = {
                        "总结标题": summary_title_display,
                        "生成时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "文档数量": doc_count,
                        "文档列表": doc_list,
                        "总结内容": st.session_state.summary
                    }
                    st.download_button(
                        label="📊 保存为JSON",
                        data=json.dumps(json_data, ensure_ascii=False, indent=2),
                        file_name=f"知识库总结_{timestamp}.json",
                        mime="application/json",
                        use_container_width=True
                    )
                
                # 自动保存到本地（可选）
                save_dir = os.path.join(".", "saved_reports")
                os.makedirs(save_dir, exist_ok=True)
                
                if st.checkbox("💾 同时自动保存到本地", value=False, key="auto_save_summary"):
                    try:
                        # 保存TXT版本（如果文件已存在则覆盖）
                        txt_path = os.path.join(save_dir, f"知识库总结_{timestamp}.txt")
                        with open(txt_path, 'w', encoding='utf-8') as f:
                            f.write(st.session_state.summary)
                        
                        # 保存Markdown版本（如果文件已存在则覆盖）
                        md_path = os.path.join(save_dir, f"知识库总结_{timestamp}.md")
                        with open(md_path, 'w', encoding='utf-8') as f:
                            f.write(md_summary)
                        
                        # 检查文件是否已存在（用于提示用户）
                        file_existed = os.path.exists(txt_path) or os.path.exists(md_path)
                        if file_existed:
                            st.success(f"✅ 报告已更新到: {save_dir}（已覆盖同名文件）")
                        else:
                            st.success(f"✅ 报告已保存到: {save_dir}")
                        st.info(f"📁 文件路径:\n- {txt_path}\n- {md_path}")
                    except Exception as e:
                        st.error(f"保存失败: {str(e)}")
    
    with col2:
        st.header("🤖 智能问答")
        
        # 聊天历史
        if st.session_state.chat_history:
            with st.expander("🗣️ 对话历史", expanded=False):
                # 保存对话历史按钮
                col_history1, col_history2 = st.columns([3, 1])
                with col_history2:
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    
                    # 生成对话历史文本
                    history_text = f"# 对话历史记录\n\n生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n文档数量: {len(st.session_state.docs)}\n\n---\n\n"
                    for i, (q, a) in enumerate(st.session_state.chat_history, 1):
                        history_text += f"## 对话 {i}\n\n**问题:** {q}\n\n**回答:**\n{a}\n\n---\n\n"
                    
                    st.download_button(
                        label="💾 保存对话",
                        data=history_text,
                        file_name=f"对话历史_{timestamp}.md",
                        mime="text/markdown",
                        use_container_width=True
                    )
                
                # 显示最近5条对话
                for i, (q, a) in enumerate(st.session_state.chat_history[-5:], start=len(st.session_state.chat_history)-4):  # 显示最近5条
                    st.markdown(f"**Q{i}:** {q}")
                    st.markdown(f"**A{i}:** {a}")
                    st.markdown("---")
        
        # 问题输入
        question = st.text_area(
            "输入您的问题",
            placeholder="例如：总结一下文档的主要内容是什么？或者：从这些文档中找出关于XX的信息。",
            height=100,
            key="question_input"
        )
        
        col_a, col_b, col_c = st.columns([1, 1, 2])
        
        with col_a:
            search_clicked = st.button("🔍 搜索答案", type="primary", use_container_width=True)
        
        with col_b:
            if st.button("🧹 清空对话", use_container_width=True):
                # 清空前询问是否保存
                if st.session_state.chat_history:
                    st.warning("⚠️ 清空对话前建议先保存对话历史！")
                st.session_state.chat_history = []
                st.rerun()
        
        with col_c:
            if st.button("💡 示例问题", use_container_width=True):
                examples = [
                    "总结所有文档的核心要点",
                    "提取文档中的关键数据",
                    "列出所有提到的重要日期",
                    "各文档之间的关联是什么？"
                ]
                st.session_state.question_input = st.session_state.get("question_input", "") + examples[0]
        
        # 处理搜索答案的逻辑（移到列布局外，使内容占据全宽）
        if search_clicked:
            # 清除高级功能显示状态
            st.session_state.show_data_analysis = False
            st.session_state.show_flowchart_gen = False
            st.session_state.show_gantt_gen = False
            
            if not question:
                st.warning("请输入问题")
            elif not api_key:
                st.error("请输入DeepSeek API密钥")
            elif not st.session_state.docs:
                st.error("请先加载文档")
            else:
                # 显示超时提示（全宽）
                timeout_info = st.session_state.get('api_timeout', 60)
                retry_info = st.session_state.get('api_max_retries', 3)
                st.info(f"⏱️ 超时设置: {timeout_info}秒 | 重试次数: {retry_info}次 | 如遇超时可在侧边栏调整")
                
                with st.spinner(f"正在思考...（超时时间: {timeout_info}秒）"):
                    answer = answer_with_deepseek(
                        question, 
                        st.session_state.vectorstore, 
                        st.session_state.docs, 
                        api_key
                    )
                    
                    # 保存到历史
                    st.session_state.chat_history.append((question, answer))
                    
                    # 显示答案（全宽）
                    st.markdown("### 💡 答案")
                    st.write(answer)
                    
                    # 保存单个问答
                    col_save_qa1, col_save_qa2 = st.columns(2)
                    timestamp_qa = datetime.now().strftime("%Y%m%d_%H%M%S")
                    
                    # 生成问答内容
                    qa_text = f"# 问答记录\n\n生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n## 问题\n\n{question}\n\n## 回答\n\n{answer}\n"
                    
                    with col_save_qa1:
                        st.download_button(
                            label="💾 保存此问答",
                            data=qa_text,
                            file_name=f"问答_{timestamp_qa}.md",
                            mime="text/markdown",
                            use_container_width=True
                        )
                    
                    with col_save_qa2:
                        # 自动保存选项
                        if st.checkbox("自动保存", key=f"auto_save_{timestamp_qa}", value=False):
                            save_dir_qa = os.path.join(".", "saved_qa")
                            os.makedirs(save_dir_qa, exist_ok=True)
                            try:
                                qa_path = os.path.join(save_dir_qa, f"问答_{timestamp_qa}.md")
                                with open(qa_path, 'w', encoding='utf-8') as f:
                                    f.write(qa_text)
                                st.success(f"✅ 已保存到: {save_dir_qa}")
                            except Exception as e:
                                st.error(f"保存失败: {str(e)}")
                    
                    # 如果答案包含错误信息，提供解决建议（不自动展开）
                    if "超时" in answer or "连接" in answer or "网络" in answer:
                        with st.expander("💡 网络问题解决建议", expanded=False):
                            st.markdown("""
                            **如果遇到超时或连接问题，可以尝试：**
                            1. 📈 **增加超时时间**：在侧边栏"高级设置"中增加超时时间（建议120-180秒）
                            2. 🔄 **增加重试次数**：在侧边栏"高级设置"中增加重试次数（建议4-5次）
                            3. 🌐 **检查网络**：确保网络连接稳定，可以访问 api.deepseek.com
                            4. 📝 **减少文档内容**：如果文档很大，尝试减少加载的文档数量
                            5. ⏰ **稍后重试**：可能是DeepSeek服务器繁忙，稍后再试
                            """)
                    
                    # 显示检索来源
                    if st.session_state.vectorstore:
                        with st.expander("查看参考来源"):
                            similar_docs = search_similar_documents(
                                st.session_state.vectorstore, 
                                question
                            )
                            for i, (content, source) in enumerate(similar_docs[:3], 1):
                                st.markdown(f"**来源 {i} - {source}**")
                                st.caption(content[:300] + "...")
        
        # 高级功能（在col2内，确保布局正确）
        st.markdown("---")
        st.subheader("🎯 高级功能")
        
        # 数据分析按钮（语义搜索功能已移除，因为与"查看参考来源"功能重复）
        data_analysis_clicked = st.button("📊 数据分析", use_container_width=True, key="data_analysis_btn")
        
        # 初始化session_state
        if 'show_data_analysis' not in st.session_state:
            st.session_state.show_data_analysis = False
        
        if data_analysis_clicked:
            # 先关闭其他功能
            st.session_state.show_flowchart_gen = False
            st.session_state.show_gantt_gen = False
            st.session_state.show_data_analysis = True
            st.rerun()  # 立即刷新页面以确保状态生效
        
        # 处理数据分析（在col2内，使用容器组织结果）
        if st.session_state.show_data_analysis:
            if not st.session_state.docs:
                st.warning("请先加载文档")
            elif not api_key:
                st.error("请输入DeepSeek API密钥")
            else:
                # 在弹窗中显示模版选择和管理
                with st.expander("📝 数据分析配置", expanded=True):
                    # 初始化session state
                    if 'selected_analysis_template' not in st.session_state:
                        st.session_state.selected_analysis_template = "default"
                    
                    # 加载模版列表
                    analysis_templates = load_templates("analysis")
                    analysis_template_options = {f"{t['name']} ({tid})": tid for tid, t in analysis_templates.items()}
                    
                    col_analysis_template1, col_analysis_template1_btn, col_analysis_template2, col_analysis_template3 = st.columns([3, 1, 1, 1])
                    with col_analysis_template1:
                        selected_analysis_template_display = st.selectbox(
                            "选择Prompt模版",
                            options=list(analysis_template_options.keys()),
                            index=list(analysis_template_options.values()).index(st.session_state.selected_analysis_template) if st.session_state.selected_analysis_template in analysis_template_options.values() else 0,
                            help="选择用于数据分析的Prompt模版",
                            key="analysis_template_select"
                        )
                        st.session_state.selected_analysis_template = analysis_template_options[selected_analysis_template_display]
                    
                    with col_analysis_template1_btn:
                        # 刷新按钮
                        st.markdown("<div style='height: 27px;'></div>", unsafe_allow_html=True)
                        if st.button("🔄 刷新", use_container_width=True, key="refresh_analysis_template_btn", help="刷新模板列表以获取最新模板"):
                            # 重新加载模板
                            analysis_templates = load_templates("analysis")
                            analysis_template_options = {f"{t['name']} ({tid})": tid for tid, t in analysis_templates.items()}
                            # 如果当前选中的模板不存在，重置为default
                            if st.session_state.selected_analysis_template not in analysis_template_options.values():
                                st.session_state.selected_analysis_template = "default"
                            st.rerun()
                    
                    with col_analysis_template2:
                        # 添加占位符以对齐selectbox的label和help icon
                        st.markdown("<div style='height: 27px;'></div>", unsafe_allow_html=True)
                        if st.button("👁️ 预览", use_container_width=True, key="preview_analysis_template"):
                            st.session_state.show_analysis_template_preview = True
                    
                    with col_analysis_template3:
                        # 添加占位符以对齐selectbox的label和help icon
                        st.markdown("<div style='height: 27px;'></div>", unsafe_allow_html=True)
                        # 检查选中的模版是否可以删除（默认模版不可删除）
                        can_delete_analysis = not is_default_template("analysis", st.session_state.selected_analysis_template)
                        if st.button("🗑️ 删除", use_container_width=True, key="delete_analysis_template_btn", disabled=not can_delete_analysis):
                            if delete_template("analysis", st.session_state.selected_analysis_template):
                                st.success("✅ 模版已删除")
                                st.session_state.selected_analysis_template = "default"
                                st.rerun()
                    
                    # 模版预览
                    if st.session_state.get('show_analysis_template_preview', False):
                        template_data = get_template("analysis", st.session_state.selected_analysis_template)
                        if template_data:
                            st.markdown("**模版预览**")
                            st.markdown(f"**模版名称**: {template_data.get('name', '')}")
                            st.markdown(f"**模版描述**: {template_data.get('description', '')}")
                            st.markdown("**模版内容**:")
                            st.code(template_data.get('template', ''), language='text')
                            if st.button("关闭预览", key="close_preview_analysis"):
                                st.session_state.show_analysis_template_preview = False
                    
                    # 模版管理（仅保留创建/编辑功能）
                    with st.expander("⚙️ 模版管理", expanded=False):
                        st.markdown("**创建/编辑模版**")
                        new_analysis_template_name = st.text_input("模版名称", key="new_analysis_template_name")
                        new_analysis_template_desc = st.text_input("模版描述", key="new_analysis_template_desc")
                        new_analysis_template_content = st.text_area(
                            "模版内容（使用 {doc_info} 作为文档信息占位符）",
                            height=200,
                            key="new_analysis_template_content",
                            help="示例：请分析以下文档集合：\n\n文档信息：\n{doc_info}\n\n分析："
                        )
                        if st.button("💾 保存模版", key="save_analysis_template"):
                            if new_analysis_template_name and new_analysis_template_content:
                                if save_template("analysis", "", new_analysis_template_name, new_analysis_template_desc, new_analysis_template_content):
                                    st.success("✅ 模版已保存")
                                    st.rerun()
                            else:
                                st.warning("请输入模版名称和内容")
                    
                    # 执行分析按钮
                    run_analysis_clicked = st.button("🚀 执行数据分析", type="primary", use_container_width=True, key="run_analysis_btn")
                    
                    if run_analysis_clicked:
                        with st.spinner("正在分析文档..."):
                            # 加载选中的模版
                            template_data = get_template("analysis", st.session_state.selected_analysis_template)
                            
                            # 准备文档信息
                            doc_info = chr(10).join([f'{name}: {len(str(data["content"]))} 字符' for name, data in st.session_state.docs.items()])
                            
                            if template_data:
                                template_str = template_data.get("template", "")
                                # 替换模版中的占位符
                                prompt = template_str.format(doc_info=doc_info)
                            else:
                                # 如果模版不存在，使用默认模版
                                prompt = f"""请分析以下文档集合，提供数据分析:

文档信息：
{doc_info}

请提供：
1. 文档内容分布分析
2. 潜在的数据模式和趋势
3. 建议的数据可视化方式"""
                            
                            analysis = query_deepseek(prompt, api_key)
                            st.session_state.analysis_result = analysis
                            st.session_state.analysis_template_name = template_data.get('name', '默认模版') if template_data else '默认模版'
                
                # 显示分析结果
                if 'analysis_result' in st.session_state and st.session_state.analysis_result:
                    st.markdown("---")
                    st.markdown(f"### 📊 数据分析结果（使用模版：{st.session_state.analysis_template_name}）")
                    st.write(st.session_state.analysis_result)
        
        # 流程图生成功能
        flowchart_gen_clicked = st.button("📐 制作流程图文件", use_container_width=True, key="flowchart_gen_btn")
        
        # 初始化session_state
        if 'show_flowchart_gen' not in st.session_state:
            st.session_state.show_flowchart_gen = False
        
        if flowchart_gen_clicked:
            # 先关闭其他功能
            st.session_state.show_data_analysis = False
            st.session_state.show_gantt_gen = False
            st.session_state.show_flowchart_gen = True
            st.rerun()  # 立即刷新页面以确保状态生效
        
        # 处理流程图生成（在col2内）
        if st.session_state.show_flowchart_gen:
            with st.expander("📐 制作流程图文件", expanded=True):
                st.markdown("""
                **使用说明：**
                1. 从左侧"批量总结"区域的总结报告中复制流程图部分的文本
                2. 将文本粘贴到下方的文本框中
                3. 点击"生成流程图"按钮
                4. 系统会自动生成并下载 draw.io 格式的流程图文件
                
                **流程图格式要求：**
                - 主要阶段使用方括号 `[]` 包裹，通过向下箭头 `↓` 连接
                - 子任务通过向右箭头 `→` 连接
                - 支持缩进表示层级关系
                """)
                
                # 文本输入区域
                flowchart_text = st.text_area(
                    "流程图文本",
                    height=300,
                    placeholder="请粘贴流程图文本，例如：\n\n[阶段一]\n↓\n[阶段二]\n→ 子任务1 → 子任务2\n↓\n[阶段三]",
                    help="从批量总结文本中复制流程图部分，粘贴到这里",
                    key="flowchart_text_input"
                )
                
                # 生成按钮
                generate_flowchart_clicked = st.button(
                    "🚀 生成流程图", 
                    type="primary", 
                    use_container_width=True,
                    key="generate_flowchart_btn"
                )
                
                if generate_flowchart_clicked:
                    if not flowchart_text or not flowchart_text.strip():
                        st.warning("请输入流程图文本")
                    else:
                        try:
                            # 导入流程图转换函数
                            from flowchart_to_drawio import convert_flowchart_to_drawio
                            
                            with st.spinner("正在生成流程图..."):
                                # 转换为 draw.io XML
                                xml_content = convert_flowchart_to_drawio(flowchart_text, None)
                                
                                # 生成文件名
                                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                file_name = f"flowchart_{timestamp}.drawio"
                                
                                # 提供下载按钮
                                st.success("✅ 流程图生成成功！")
                                st.download_button(
                                    label="📥 下载流程图文件",
                                    data=xml_content,
                                    file_name=file_name,
                                    mime="application/xml",
                                    use_container_width=True,
                                    key=f"download_flowchart_{timestamp}"
                                )
                                
                                # 显示预览提示
                                st.info("💡 提示：下载后可以使用 [draw.io](https://app.diagrams.net/) 或 [diagrams.net](https://www.diagrams.net/) 打开文件进行编辑")
                                
                        except ImportError:
                            st.error("无法导入流程图转换模块，请确保 flowchart_to_drawio.py 文件存在")
                        except Exception as e:
                            st.error(f"生成流程图时出错：{str(e)}")
                            import traceback
                            with st.expander("错误详情", expanded=False):
                                st.code(traceback.format_exc(), language='python')
        
        # 甘特图生成功能
        gantt_gen_clicked = st.button("📅 制作甘特图文件", use_container_width=True, key="gantt_gen_btn")
        
        # 初始化session_state
        if 'show_gantt_gen' not in st.session_state:
            st.session_state.show_gantt_gen = False
        
        if gantt_gen_clicked:
            # 先关闭其他功能
            st.session_state.show_data_analysis = False
            st.session_state.show_flowchart_gen = False
            st.session_state.show_gantt_gen = True
            st.rerun()  # 立即刷新页面以确保状态生效
        
        # 处理甘特图生成（在col2内）
        if st.session_state.show_gantt_gen:
            with st.expander("📅 制作甘特图文件", expanded=True):
                st.markdown("""
                **使用说明：**
                1. 准备项目进度甘特图表数据（表格格式）
                2. 将表格数据粘贴到下方的文本框中
                3. 点击"生成甘特图"按钮
                4. 系统会自动生成并下载 draw.io 格式的甘特图文件
                
                **甘特图数据格式要求：**
                - 表格格式，使用制表符（Tab）分隔各列
                - 必须包含表头行：任务ID、任务名称、开始时间、结束时间、工期(月)、前置任务、责任方/备注
                - 时间格式：M0, M1, M1+0.5, M1.5 等
                - 任务ID支持层级结构：1, 1.1, 1.2, 2.1 等
                
                **提示：** 可以使用AI生成甘特图数据，在"批量总结"功能中选择"项目进度甘特图模版"，系统会根据文档内容自动生成符合格式要求的甘特图表数据
                """)
                
                # 文本输入区域
                gantt_text = st.text_area(
                    "甘特图表数据",
                    height=400,
                    placeholder="任务ID	任务名称	开始时间	结束时间	工期(月)	前置任务	责任方/备注\n1	项目启动	M0	M1	2		\n1.1	项目立项	M0	M0+0.5	0.5		甲方、乙方\n1.2	需求调研	M0+0.5	M1	0.5	1.1	乙方",
                    help="粘贴甘特图表数据（表格格式），支持制表符或空格分隔",
                    key="gantt_text_input"
                )
                
                # 生成按钮
                generate_gantt_clicked = st.button(
                    "🚀 生成甘特图", 
                    type="primary", 
                    use_container_width=True,
                    key="generate_gantt_btn"
                )
                
                if generate_gantt_clicked:
                    if not gantt_text or not gantt_text.strip():
                        st.warning("请输入甘特图表数据")
                    else:
                        try:
                            # 导入甘特图转换函数
                            from gantt_to_drawio import convert_gantt_to_drawio
                            
                            # 转换为 draw.io XML
                            with st.spinner("正在生成甘特图..."):
                                xml_content = convert_gantt_to_drawio(gantt_text, None)
                            
                            # spinner结束后再显示成功信息和下载按钮
                            # 生成文件名
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            file_name = f"gantt_{timestamp}.drawio"
                            
                            # 提供下载按钮
                            st.success("✅ 甘特图生成成功！")
                            st.download_button(
                                label="📥 下载甘特图文件",
                                data=xml_content,
                                file_name=file_name,
                                mime="application/xml",
                                use_container_width=True,
                                key=f"download_gantt_{timestamp}"
                            )
                            
                            # 显示预览提示
                            st.info("💡 提示：下载后可以使用 [draw.io](https://app.diagrams.net/) 或 [diagrams.net](https://www.diagrams.net/) 打开文件进行编辑")
                                
                        except ImportError:
                            st.error("无法导入甘特图转换模块，请确保 gantt_to_drawio.py 文件存在")
                        except Exception as e:
                            st.error(f"生成甘特图时出错：{str(e)}")
                            import traceback
                            with st.expander("错误详情", expanded=False):
                                st.code(traceback.format_exc(), language='python')
        
        # 显示版权信息
        show_footer()

# 简易版（无向量数据库）
def simple_main():
    """简易版本，不使用向量数据库"""
    st.set_page_config(
        page_title="简易知识库浏览器",
        page_icon="📁",
        layout="wide"
    )
    
    st.title("📁 简易知识库浏览器")
    
    # 文件上传
    uploaded_files = st.file_uploader(
        "上传文件 (支持txt, docx, pdf, excel, md, js, json)",
        type=['txt', 'docx', 'pdf', 'xlsx', 'xls', 'md', 'js', 'json'],
        accept_multiple_files=True
    )
    
    if uploaded_files:
        docs = {}
        temp_dir = tempfile.mkdtemp()
        for uploaded_file in uploaded_files:
            # 保存上传的文件到临时目录
            file_path = os.path.join(temp_dir, uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # 读取文件
            filename = uploaded_file.name
            file_ext = filename.split('.')[-1].lower()
            
            try:
                if file_ext == 'txt':
                    content = read_text_file(file_path)
                elif file_ext == 'docx':
                    content = read_docx_file(file_path)
                elif file_ext == 'pdf':
                    content = read_pdf_file(file_path)
                elif file_ext in ['xlsx', 'xls']:
                    content = read_excel_file(file_path)
                elif file_ext == 'md':
                    content = read_markdown_file(file_path)
                elif file_ext == 'js':
                    content = read_javascript_file(file_path)
                elif file_ext == 'json':
                    content = read_json_file(file_path)
                else:
                    content = f"不支持的文件类型"
                
                docs[filename] = {
                    'content': content,
                    'type': file_ext,
                    'path': file_path
                }
            except Exception as e:
                st.error(f"读取文件 {filename} 失败: {str(e)}")
        
        # 文件浏览器
        if docs:
            selected_file = st.selectbox("选择文件", list(docs.keys()))
            
            if selected_file in docs:
                content = docs[selected_file]['content']
                st.text_area("文件内容", 
                           str(content) if not isinstance(content, dict) else str(content),
                           height=400)
        
        # 简单问答
        st.markdown("---")
        
        # 尝试加载保存的 API key
        saved_api_key = load_api_key()
        default_key = saved_api_key if saved_api_key else ""
        
        col_simple_key1, col_simple_key2 = st.columns([3, 1])
        with col_simple_key1:
            api_key = st.text_input("DeepSeek API密钥", value=default_key, type="password")
        with col_simple_key2:
            if st.button("💾 保存", use_container_width=True):
                if api_key:
                    if save_api_key(api_key):
                        st.success("✅ 已保存")
                else:
                    st.warning("请输入密钥")
        
        question = st.text_input("输入问题")
        
        if st.button("获取答案") and api_key and question:
            # 合并所有文档内容
            all_content = ""
            for filename, data in docs.items():
                content = data['content']
                if isinstance(content, dict):
                    content = "\n".join([f"{k}: {v}" for k, v in content.items()])
                all_content += f"\n\n文件: {filename}\n{content}"
            
            # 调用DeepSeek
            prompt = f"""基于以下文档内容回答问题：

{all_content[:8000]}

问题：{question}

请基于文档内容回答，如果文档中没有相关信息，请明确说明。"""
            
            with st.spinner("正在思考..."):
                answer = query_deepseek(prompt, api_key)
                st.write("**答案：**", answer)
        
        # 显示版权信息
        show_footer()

if __name__ == "__main__":
    # 默认使用完整版，可以通过环境变量或命令行参数切换
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "simple":
        simple_main()
    else:
        main()